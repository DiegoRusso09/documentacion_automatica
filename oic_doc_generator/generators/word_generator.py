# =========================================================
# FILE:
# oic_doc_generator/generators/word_generator.py
# =========================================================

from io import BytesIO
from datetime import datetime
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)
from docx.oxml import parse_xml
from docx.oxml.ns import (
    nsdecls,
    qn
)

from parsers.iar_parser import (
    find_all_iar_files,
    extract_iar
)

from parsers.project_parser import (
    get_project_name,
    read_project_xml,
    build_application_map,
    get_project_root,
    get_project_version
)

from parsers.orchestration_parser import (
    get_endpoint_flows,
    generate_description
)

from parsers.connections_parser import (
    analyze_dbaas_jca,
    find_application_folder,
    build_action_description
)

from parsers.schedule_parser import (
    get_schedule_information,
    is_scheduled_integration
)

from parsers.service_parser import (
    get_request_response_schemas
)

from generators.service_design_generator import (
    add_service_design_section
)

from utils.xml_utils import (
    clean_tag,
    extract_application_from_refuri
)

from utils.mermaid_utils import (
    generate_sequence_diagram_png
)

from parsers.vb_extractor import (
    build_extraction_metadata
)

from parsers.vb_parser import (
    build_page_metadata
)

from renderers.html_builder import (
    build_complete_html
)

from renderers.screenshot_renderer import (
    render_html_to_image
)


# =========================================================
# CREATE HEADER
# =========================================================

def create_header(
    document,
    text,
    size=16
):

    p = document.add_paragraph()

    run = p.add_run(text)

    run.bold = True

    run.font.name = "Arial"

    run.font.size = Pt(size)

    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )


# =========================================================
# APPLY HEADER STYLE
# =========================================================

def apply_header_style(
    cell,
    fill="D9D9D9",
    white_text=False
):

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="{}"/>'.format(
            nsdecls('w'),
            fill
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.bold = True

            if white_text:

                run.font.color.rgb = RGBColor(
                    255,
                    255,
                    255
                )


# =========================================================
# FORMAT JSON WRAPPER
# =========================================================

def format_json_wrapper(
    elements,
    is_array=False
):

    if not elements:

        return "{}"

    lines = []

    if is_array:

        lines.append("[")
        lines.append("{")

    else:

        lines.append("{")

    for element in elements:

        name = element.get(
            "name",
            ""
        )

        data_type = element.get(
            "type",
            "string"
        )

        if name in [

            "response-wrapper",

            "request-wrapper",

            "topLevelArray"
        ]:

            continue

        lines.append(
            f'  "{name}": "{data_type}",'
        )

    if len(lines) > 1:

        lines[-1] = lines[-1].rstrip(",")

    if is_array:

        lines.append("}")
        lines.append("]")

    else:

        lines.append("}")

    return "\n".join(lines)


# =========================================================
# ADD PROCESS TABLE
# =========================================================

def add_process_table(
    document,
    rows
):

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.style = "Table Grid"

    headers = [

        "Nro. Secuencia",

        "Nombre Acción",

        "Descripción de la Acción"
    ]

    hdr = table.rows[0].cells

    for i in range(3):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )

    for row in rows:

        cells = table.add_row().cells

        cells[0].text = str(
            row["Nro. Secuencia"]
        )

        cells[1].text = (
            row["Nombre Acción"]
        )

        cells[2].text = (
            row["Descripción de la Acción"]
        )


# =========================================================
# GENERATE WORD DOCUMENT
# =========================================================

def generate_word_document(
    package_path,
    author_name,
    development_name,
    selected_components,
    visual_builder_apps,
    apex_apps,
    use_oic
):

    document = Document()

    style = document.styles['Normal']

    style.font.name = 'Arial'

    style._element.rPr.rFonts.set(

        qn('w:eastAsia'),

        'Arial'
    )

    style.font.size = Pt(10)

    # =====================================================
    # 1 CONTROL DOCUMENTO
    # =====================================================

    create_header(
        document,
        "1\tControl del Documento"
    )

    create_header(
        document,
        "1.1\tRegistro de Modificaciones"
    )

    table = document.add_table(
        rows=2,
        cols=4
    )

    table.style = "Table Grid"

    headers = [

        "Fecha",

        "Autor",

        "Versión",

        "Descripción del Cambio"
    ]

    hdr = table.rows[0].cells

    for i in range(4):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )

    today = datetime.today().strftime(
        "%d/%m/%Y"
    )

    row = table.rows[1].cells

    row[0].text = today
    row[1].text = author_name
    row[2].text = "1.0"
    row[3].text = (
        "Creación del Documento"
    )

    # =====================================================
    # 1.2 REVISORES
    # =====================================================

    create_header(
        document,
        "1.2\tRevisores"
    )

    reviewers_table = document.add_table(
        rows=5,
        cols=2
    )

    reviewers_table.style = "Table Grid"

    hdr = reviewers_table.rows[0].cells

    hdr[0].text = "Nombre"
    hdr[1].text = "Puesto"

    apply_header_style(
        hdr[0]
    )

    apply_header_style(
        hdr[1]
    )

    # =====================================================
    # 2 VISIÓN GENERAL
    # =====================================================

    create_header(
        document,
        "2\tVisión General"
    )

    document.add_paragraph(

        f"En este documento se tiene "
        f"como propósito describir "
        f"cada componente del "
        f"{development_name} y su "
        f"funcionalidad a nivel "
        f"técnico-funcional."
    )

    # =====================================================
    # 2.1
    # =====================================================

    create_header(
        document,
        "2.1\tReglas del Negocio"
    )

    for _ in range(4):

        document.add_paragraph("")

    # =====================================================
    # 2.2
    # =====================================================

    create_header(
        document,
        "2.2\tListado de Componentes"
    )

    document.add_paragraph(

        "El detalle de los "
        "componentes necesarios "
        "para el presente desarrollo "
        "se lista a continuación:"
    )

    components_table = document.add_table(
        rows=1,
        cols=2
    )

    components_table.style = "Table Grid"

    hdr = components_table.rows[0].cells

    hdr[0].text = "Componente"
    hdr[1].text = "Descripción"

    apply_header_style(
        hdr[0],
        fill="808080",
        white_text=True
    )

    apply_header_style(
        hdr[1],
        fill="808080",
        white_text=True
    )

    for component in selected_components:

        row = (
            components_table.add_row().cells
        )

        row[0].text = component
        row[1].text = ""

    # =====================================================
    # 2.3
    # =====================================================

    create_header(
        document,
        "2.3\tDiagrama de componentes"
    )

    document.add_paragraph(

        "El siguiente diagrama "
        "representa los componentes "
        "y la relación existente "
        "entre ellos."
    )

    document.add_paragraph("")
    document.add_paragraph("")

    # =====================================================
    # 2.4
    # =====================================================

    create_header(
        document,
        "2.4\tDiagrama de Secuencia"
    )

    document.add_paragraph(

        "El siguiente diagrama "
        "representa la secuencia "
        "que existe entre los "
        "componentes involucrados:"
    )

    document.add_paragraph("")
    document.add_paragraph("")

    # =====================================================
    # 2.5 OIC
    # =====================================================

    create_header(
        document,
        "2.5\tDiagrama de bloques de construcción"
    )

    integrations_data = []

    if not use_oic:

        document.add_paragraph(
            "No aplica."
        )

    else:

        iar_files = find_all_iar_files(
            package_path
        )

        for iar in iar_files:

            extracted_iar = extract_iar(
                iar
            )

            integration_name = (
                get_project_name(
                    extracted_iar
                )
            )

            applications = (
                read_project_xml(
                    extracted_iar
                )
            )

            app_map = build_application_map(
                applications
            )

            root = get_project_root(
                extracted_iar
            )

            integrations_data.append({

                "name":
                    integration_name,

                "path":
                    extracted_iar,

                "scheduled":
                    is_scheduled_integration(
                        extracted_iar
                    ),

                "version":
                    get_project_version(
                        extracted_iar
                    )
            })

            p = document.add_paragraph()

            run = p.add_run(
                f"• {integration_name}"
            )

            run.bold = True

            endpoint_flows = (
                get_endpoint_flows(
                    root,
                    app_map
                )
            )

            # =============================================
            # ENDPOINTS
            # =============================================

            for endpoint, flow in (

                endpoint_flows.items()
            ):

                description = (
                    generate_description(

                        endpoint,

                        flow,

                        applications,

                        extracted_iar,

                        root
                    )
                )

                document.add_paragraph(
                    description
                )

                document.add_paragraph(

                    "A continuación, "
                    "se describe cada "
                    "una de las acciones "
                    "del flujo de integración:"
                )

                rows = []

                seq = 1

                if is_scheduled_integration(
                    extracted_iar
                ):

                    start_desc = (
                        "Integración programada."
                    )

                else:

                    start_desc = (
                        "De tipo REST, endpoint."
                    )

                rows.append({

                    "Nro. Secuencia":
                        seq,

                    "Nombre Acción":
                        endpoint,

                    "Descripción de la Acción":
                        start_desc
                })

                seq += 1

                # =========================================
                # FLOW ACTIONS
                # =========================================

                for elem in flow.iter():

                    try:

                        tag = clean_tag(
                            elem.tag
                        )

                    except:

                        continue

                    # =====================================
                    # INVOKE
                    # =====================================

                    if tag == "invoke":

                        refuri = elem.attrib.get(
                            "refUri",
                            ""
                        )

                        app_ref = (
                            extract_application_from_refuri(
                                refuri
                            )
                        )

                        if (
                            app_ref
                            not in app_map
                        ):

                            continue

                        app = app_map[
                            app_ref
                        ]

                        folder = (
                            find_application_folder(
                                extracted_iar,
                                app_ref
                            )
                        )

                        dbaas = (
                            analyze_dbaas_jca(
                                folder
                            )
                        )

                        desc = build_action_description(
                            app,
                            dbaas
                        )

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                app["Invoke"],

                            "Descripción de la Acción":
                                desc
                        })

                        seq += 1

                    # =====================================
                    # ASSIGNMENT
                    # =====================================

                    elif tag in [

                        "assign",

                        "assignment"
                    ]:

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "Assignment",

                            "Descripción de la Acción":
                                "Se realiza un Assignment."
                        })

                        seq += 1

                    # =====================================
                    # SWITCH
                    # =====================================

                    elif tag in [

                        "switch",

                        "router"
                    ]:

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "Switch",

                            "Descripción de la Acción":
                                "Se realiza un Switch."
                        })

                        seq += 1

                    # =====================================
                    # STAGE FILE
                    # =====================================

                    elif tag == "stageFile":

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "stageFile",

                            "Descripción de la Acción":
                                "Se realiza un stageFile."
                        })

                        seq += 1

                    # =====================================
                    # FOR
                    # =====================================

                    elif tag == "for":

                        refuri = elem.attrib.get(
                            "refUri",
                            ""
                        )

                        processor_name = ""
                        iteration_name = ""
                        parallel = "false"

                        for proc in flow.iter():

                            try:

                                proc_tag = clean_tag(
                                    proc.tag
                                )

                            except:

                                continue

                            if proc_tag != "processor":

                                continue

                            if proc.attrib.get(
                                "name",
                                ""
                            ) != refuri:

                                continue

                            for child in proc.iter():

                                try:

                                    child_tag = clean_tag(
                                        child.tag
                                    )

                                except:

                                    continue

                                if child_tag.lower() == "processorname":

                                    if child.text:

                                        processor_name = (
                                            child.text.strip()
                                        )

                                elif child_tag.lower() == "subrole":

                                    if child.text:

                                        iteration_name = (
                                            child.text.strip()
                                        )

                                elif child_tag.lower() == "property":

                                    if (
                                        child.attrib.get(
                                            "name",
                                            ""
                                        )
                                        ==
                                        "parallel"
                                    ):

                                        parallel = child.attrib.get(
                                            "value",
                                            "false"
                                        )

                            break

                        mode = (
                            "paralela"
                            if parallel.lower() == "true"
                            else "secuencial"
                        )

                        if iteration_name:

                            description = (

                                f"Ahora comienza un For "
                                f"que itera sobre "
                                f"{iteration_name} "
                                f"de manera {mode}."
                            )

                        else:

                            description = (
                                "Se realiza un For."
                            )

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                processor_name or "For",

                            "Descripción de la Acción":
                                description
                        })

                        seq += 1

                # =========================================
                # TABLE
                # =========================================

                add_process_table(
                    document,
                    rows
                )

                # =========================================
                # DIAGRAM
                # =========================================

                document.add_paragraph(
                    "Diagrama de Secuencia del Endpoint:"
                )

                integration_type = "REST"

                if is_scheduled_integration(
                    extracted_iar
                ):

                    integration_type = "scheduled"

                img_path = (
                    generate_sequence_diagram_png(

                        endpoint,

                        rows,

                        integration_type
                    )
                )

                if img_path and os.path.exists(
                    img_path
                ):

                    document.add_picture(

                        img_path,

                        width=Inches(6.5)
                    )

                # =========================================
                # FREQUENCY
                # =========================================

                if is_scheduled_integration(
                    extracted_iar
                ):

                    schedule_info = (
                        get_schedule_information(
                            extracted_iar
                        )
                    )

                    frequency = (
                        schedule_info.get(
                            "ical_expression",
                            "No definida"
                        )
                    )

                else:

                    frequency = "No aplica"

                document.add_paragraph(

                    f"Frecuencia de Ejecución\t: "
                    f"{frequency}"
                )

                document.add_paragraph("")
                document.add_paragraph("")

    # =====================================================
    # 3 DISEÑO DE PANTALLA
    # =====================================================

    create_header(
        document,
        "3\tDiseño de Pantalla"
    )

    screen_counter = 1

    # =====================================================
    # VALIDATION
    # =====================================================

    if not visual_builder_apps:

        document.add_paragraph(
            "No se encontraron aplicaciones Visual Builder."
        )

    else:

        # =================================================
        # NORMALIZE
        # =================================================

        if not isinstance(
            visual_builder_apps,
            list
        ):

            visual_builder_apps = [
                visual_builder_apps
            ]

        # =================================================
        # APPLICATIONS
        # =================================================

        for vb_zip in visual_builder_apps:

            try:

                # =========================================
                # EXTRACT
                # =========================================

                extraction_metadata = (
                    build_extraction_metadata(
                        vb_zip
                    )
                )

                # =========================================
                # PARSE PAGES
                # =========================================

                vb_metadata = (
                    build_page_metadata(
                        extraction_metadata
                    )
                )

                shell_html = vb_metadata.get(
                    "shell_html",
                    ""
                )

                app_css = vb_metadata.get(
                    "app_css",
                    ""
                )

                resources_path = vb_metadata.get(
                    "resources_path"
                )

                pages = vb_metadata.get(
                    "pages",
                    []
                )

                # =========================================
                # NO PAGES
                # =========================================

                if not pages:

                    document.add_paragraph(
                        "No se encontraron páginas Visual Builder."
                    )

                    continue

                # =========================================
                # LOOP PAGES
                # =========================================

                for page in pages:

                    page_name = page.get(
                        "name",
                        f"Pantalla {screen_counter}"
                    )

                    page_html = page.get(
                        "html",
                        ""
                    )

                    create_header(

                        document,

                        f"3.{screen_counter}\t{page_name}",

                        size=14
                    )

                    try:

                        # =================================
                        # BUILD HTML
                        # =================================

                        complete_html = (
                            build_complete_html(

                                shell_html=
                                    shell_html,

                                page_html=
                                    page_html,

                                app_css=
                                    app_css
                            )
                        )

                        # =================================
                        # RENDER IMAGE
                        # =================================

                        image_path = (
                            render_html_to_image(

                                html_content=
                                    complete_html,

                                resources_path=
                                    resources_path
                            )
                        )

                        # =================================
                        # INSERT IMAGE
                        # =================================

                        if (

                            image_path

                            and

                            os.path.exists(
                                image_path
                            )
                        ):

                            document.add_picture(

                                image_path,

                                width=Inches(6.5)
                            )

                        else:

                            document.add_paragraph(
                                "No fue posible generar la imagen."
                            )

                    except Exception as e:

                        document.add_paragraph(
                            f"Error renderizando pantalla: {str(e)}"
                        )

                    document.add_paragraph("")

                    screen_counter += 1

            except Exception as e:

                document.add_paragraph(
                    f"Error procesando Visual Builder: {str(e)}"
                )



    # =====================================================
    # 4 DISEÑO DE SERVICIOS
    # =====================================================

    create_header(
        document,
        "4\tDiseño de Interface"
    )

    create_header(
        document,
        "4.1\tDiseño de Servicio",
        size=14
    )

    if not use_oic:

        document.add_paragraph(
            "No aplica."
        )

    else:

        if not integrations_data:

            document.add_paragraph(
                "No se encontraron integraciones."
            )

        else:

            for integration in integrations_data:

                if integration.get(
                    "scheduled",
                    False
                ):

                    continue
                try:

                    add_service_design_section(

                        document,

                        integration.get(
                            "path"
                        ),

                        integration.get(
                            "name"
                        ),

                        integration.get(
                            "version",
                            "01.00.0000"
                        )
                    )

                    document.add_paragraph("")
                    document.add_paragraph("")

                except Exception as e:

                    document.add_paragraph(

                        f"Error generando diseño "
                        f"de servicio: {str(e)}"
                    )


    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output