# =========================================================
# FILE: service_design_generator.py
# =========================================================

from docx.shared import Pt

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.enum.table import (
    WD_TABLE_ALIGNMENT
)

from docx.oxml import parse_xml

from docx.oxml.ns import (
    nsdecls
)

from parsers.service_parser import (
    get_service_design
)


# =========================================================
# HEADER STYLE
# =========================================================

def style_header_cell(cell):

    for paragraph in cell.paragraphs:

        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        if paragraph.runs:

            paragraph.runs[0].bold = True

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="D9D9D9"/>'.format(
            nsdecls('w')
        )

    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )


# =========================================================
# CREATE REQUEST/RESPONSE TABLE
# =========================================================

def create_service_table(
    document,
    request_fields,
    response_fields
):

    table = document.add_table(
        rows=4,
        cols=1
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    # =====================================================
    # REQUEST HEADER
    # =====================================================

    request_header = table.cell(0, 0)

    request_header.text = "REQUEST"

    style_header_cell(
        request_header
    )

    # =====================================================
    # REQUEST BODY
    # =====================================================

    request_body = table.cell(1, 0)

    if request_fields:

        lines = []

        for field in request_fields:

            lines.append(

                f"{field['name']} "
                f"({field['type']})"

            )

        request_body.text = (
            "\n".join(lines)
        )

    else:

        request_body.text = (
            "No aplica."
        )

    # =====================================================
    # RESPONSE HEADER
    # =====================================================

    response_header = table.cell(2, 0)

    response_header.text = "RESPONSE"

    style_header_cell(
        response_header
    )

    # =====================================================
    # RESPONSE BODY
    # =====================================================

    response_body = table.cell(3, 0)

    if response_fields:

        lines = []

        for field in response_fields:

            lines.append(

                f"{field['name']} "
                f"({field['type']})"

            )

        response_body.text = (
            "\n".join(lines)
        )

    else:

        response_body.text = (
            "No aplica."
        )


# =========================================================
# ADD SERVICE DESIGN SECTION
# =========================================================

def add_service_design_section(
    document,
    integration_path,
    integration_name,
    version
):

    services = get_service_design(
        integration_path
    )

    # =====================================================
    # TITLE
    # =====================================================

    title = document.add_paragraph()

    run = title.add_run(
        "4.1\tDiseño de Servicio"
    )

    run.bold = True

    run.font.size = Pt(15)

    # =====================================================
    # SERVICES
    # =====================================================

    if not services:

        document.add_paragraph(
            "No se encontraron diseños de servicio."
        )

        return

    for service in services:

        # =================================================
        # INTEGRATION TITLE
        # =================================================

        document.add_paragraph(

            f"• {integration_name} "
            f"| {version}"

        )

        # =================================================
        # ENDPOINT INFO
        # =================================================

        endpoint_name = service.get(
            "endpoint_name",
            "UNKNOWN"
        )

        has_request = bool(
            service.get(
                "request_fields",
                []
            )
        )

        has_response = bool(
            service.get(
                "response_fields",
                []
            )
        )

        endpoint_line = endpoint_name

        if (
            has_request
            and
            has_response
        ):

            endpoint_line += (
                " + Request/Response"
            )

        elif has_request:

            endpoint_line += (
                " + Request"
            )

        elif has_response:

            endpoint_line += (
                " + Response"
            )

        document.add_paragraph(
            endpoint_line
        )

        # =================================================
        # DATA TYPE
        # =================================================

        document.add_paragraph(
            "Tipo de Dato\t: JSON"
        )

        # =================================================
        # TABLE
        # =================================================

        create_service_table(

            document,

            service.get(
                "request_fields",
                []
            ),

            service.get(
                "response_fields",
                []
            )

        )

        document.add_paragraph("")