# =========================================================
# FILE:
# oic_doc_generator/generators/database_design_generator.py
# PART 1
# =========================================================

import os

from docx.shared import (
    Inches
)

from generators.erd_generator import (
    generate_erd_diagram
)

from utils.word_utils import (
    create_header,
    add_description_box
)


# =========================================================
# ADD TABLE HEADER
# =========================================================

def add_table_title(
    document,
    title
):

    p = document.add_paragraph()

    run = p.add_run(
        f"•\t{title}"
    )

    run.bold = True


# =========================================================
# ADD TABLE DEFINITION
# =========================================================

def add_table_definition(
    document,
    table_metadata
):

    columns = table_metadata.get(
        "columns",
        []
    )

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = (
        "Nombre de Columna"
    )

    header[1].text = (
        "Tipo de Dato"
    )

    header[2].text = (
        "Tamaño"
    )

    header[3].text = (
        "Descripción"
    )

    for column in columns:

        row = table.add_row().cells

        row[0].text = str(
            column.get(
                "column_name",
                ""
            )
        )

        row[1].text = str(
            column.get(
                "data_type",
                ""
            )
        )

        row[2].text = str(
            column.get(
                "size",
                ""
            )
        )

        row[3].text = ""

    document.add_paragraph("")


# =========================================================
# ADD ERD IMAGE
# =========================================================

def add_erd_image(
    document,
    tables
):

    try:

        image_path = generate_erd_diagram(
            tables
        )

        if (
            image_path
            and
            os.path.exists(
                image_path
            )
        ):

            document.add_picture(

                image_path,

                width=Inches(6.5)
            )

            document.add_paragraph("")

    except Exception as e:

        document.add_paragraph(

            f"Error DER: {str(e)}"
        )


# =========================================================
# SECTION 7.1 TABLES
# =========================================================

def add_tables_section(
    document,
    database_metadata
):

    create_header(
        document,
        "7.1\tTablas"
    )

    add_description_box(

        document,

        "Se presenta una imagen "
        "del modelo lógico de "
        "entidad relación para "
        "las tablas personalizadas "
        "que participan en el desarrollo."
    )

    document.add_paragraph("")

    # =====================================================
    # ERD
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        "•\tModelo Entidad Relación"
    )

    run.bold = True

    tables = database_metadata.get(
        "tables",
        []
    )

    add_erd_image(
        document,
        tables
    )

    # =====================================================
    # TABLE DEFINITIONS
    # =====================================================

    for table_metadata in tables:

        add_table_title(

            document,

            table_metadata.get(
                "table_name",
                ""
            )
        )

        add_table_definition(

            document,

            table_metadata
        )


# =========================================================
# CONTINUATION
# FILE:
# oic_doc_generator/generators/database_design_generator.py
# PART 2
# =========================================================

from utils.sql_exporter import (
    export_database_sql
)


# =========================================================
# SECTION 7.2 SEQUENCES
# =========================================================

def add_sequences_section(
    document,
    database_metadata,
    database_export_info
):

    create_header(
        document,
        "7.2\tSecuencias"
    )

    sequences = database_metadata.get(
        "sequences",
        []
    )

    if not sequences:

        document.add_paragraph(
            "No se encontraron secuencias."
        )

        return

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = "Nombre"
    header[1].text = "Número de Inicio"
    header[2].text = "Descripción"
    header[3].text = "Script"

    sequence_files = (
        database_export_info.get(
            "sequence_files",
            {}
        )
        if database_export_info
        else {}
    )

    for sequence in sequences:

        row = table.add_row().cells

        sequence_name = sequence.get(
            "sequence_name",
            ""
        )

        row[0].text = sequence_name

        row[1].text = str(
            sequence.get(
                "start_with",
                ""
            )
        )

        row[2].text = ""

        script_path = sequence_files.get(
            sequence_name,
            ""
        )

        if script_path:

            relative_path = os.path.join(
                "SQL",
                "Secuencias",
                os.path.basename(
                    script_path
                )
            )

            row[3].text = relative_path

        else:

            row[3].text = ""

    document.add_paragraph("")


# =========================================================
# PACKAGE TABLE
# =========================================================

def add_package_table(
    document,
    packages,
    package_files
):

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = (
        "Nombre de Paquete"
    )

    header[1].text = (
        "Descripción"
    )

    header[2].text = (
        "Package Spec"
    )

    header[3].text = (
        "Package Body"
    )

    for package in packages:

        row = table.add_row().cells

        package_name = package.get(
            "package_name",
            ""
        )

        row[0].text = package_name

        row[1].text = ""

        package_info = package_files.get(
            package_name,
            {}
        )

        spec_path = package_info.get(
            "spec",
            ""
        )

        body_path = package_info.get(
            "body",
            ""
        )

        if spec_path:

            row[2].text = os.path.join(
                "SQL",
                "Packages",
                os.path.basename(
                    spec_path
                )
            )

        if body_path:

            row[3].text = os.path.join(
                "SQL",
                "Packages",
                os.path.basename(
                    body_path
                )
            )


# =========================================================
# SECTION 7.3 PACKAGES
# =========================================================

def add_packages_section(
    document,
    database_metadata,
    database_export_info
):

    create_header(
        document,
        "7.3\tPaquetes"
    )

    packages = database_metadata.get(
        "packages",
        []
    )

    if not packages:

        document.add_paragraph(
            "No se encontraron paquetes."
        )

        return

    package_files = (
        database_export_info.get(
            "package_files",
            {}
        )
        if database_export_info
        else {}
    )

    add_package_table(

        document,

        packages,

        package_files
    )

    document.add_paragraph("")


# =========================================================
# MEMBERS TABLE
# =========================================================

def add_members_table(
    document,
    package
):

    package_name = package.get(
        "package_name",
        ""
    )

    members = package.get(
        "members",
        []
    )

    p = document.add_paragraph()

    run = p.add_run(
        f"•\t{package_name}"
    )

    run.bold = True

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = (
        "Función o Procedimiento"
    )

    header[1].text = (
        "Tipo"
    )

    header[2].text = (
        "Descripción"
    )

    for member in members:

        row = table.add_row().cells

        row[0].text = str(
            member.get(
                "name",
                ""
            )
        )

        row[1].text = str(
            member.get(
                "type",
                ""
            )
        )

        row[2].text = ""

    document.add_paragraph("")


# =========================================================
# SECTION 7.4 FUNCTIONS
# =========================================================

def add_functions_section(
    document,
    database_metadata
):

    create_header(
        document,
        "7.4\tFunciones y Procedimientos"
    )

    packages = database_metadata.get(
        "packages",
        []
    )

    if not packages:

        document.add_paragraph(
            "No se encontraron paquetes."
        )

        return

    for package in packages:

        add_members_table(
            document,
            package
        )


# =========================================================
# MAIN SECTION
# =========================================================

def add_database_design_section(
    document,
    database_metadata,
    database_export_info=None
):

    create_header(
        document,
        "7\tDiseño de Base de Datos"
    )

    add_description_box(

        document,

        "En esta sección de describen "
        "los objetos de base de datos "
        "que utiliza el desarrollo, "
        "estos objetos son aquellos "
        "que han sido creados como "
        "parte del desarrollo en la "
        "base de datos personalizada."
    )

    document.add_paragraph("")

    add_tables_section(
        document,
        database_metadata
    )

    add_sequences_section(
        document,
        database_metadata,
        database_export_info
    )

    add_packages_section(
        document,
        database_metadata,
        database_export_info
    )

    add_functions_section(
        document,
        database_metadata
    )