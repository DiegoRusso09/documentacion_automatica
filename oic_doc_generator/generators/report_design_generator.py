# =========================================================
# FILE:
# oic_doc_generator/generators/report_design_generator.py
# =========================================================

from docx.shared import (
    Inches,
    Pt,
    Cm,
    RGBColor
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.enum.table import (
    WD_TABLE_ALIGNMENT
)

from docx.oxml import parse_xml

from docx.oxml.ns import (
    nsdecls
)

from docx.oxml.ns import (
    nsdecls
)

# =========================================================
# APPLY HEADER STYLE
# =========================================================

def apply_header_style(
    cell,
    fill="D9D9D9",
    white_text=False
):

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="{}"/>'.format(
            nsdecls('w'),
            fill
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.bold = True

            if white_text:

                run.font.color.rgb = (
                    RGBColor(
                        255,
                        255,
                        255
                    )
                )


# =========================================================
# CREATE HEADER
# =========================================================

def create_header(
    document,
    text,
    size=16
):

    # =====================================================
    # HD1
    # =====================================================

    if text.startswith("5\t"):

        p = document.add_paragraph(
            style="HD1"
        )

    # =====================================================
    # HD2
    # =====================================================

    else:

        p = document.add_paragraph(
            style="HD2"
        )

    run = p.add_run(text)

    run.bold = True

    run.font.name = "Arial"

    run.font.size = Pt(size)


# =========================================================
# ADD DESCRIPTION BOX
# =========================================================

def add_description_box(
    document,
    text
):

    table = document.add_table(
        rows=1,
        cols=1
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    table.autofit = False

    cell = table.cell(0,0)

    cell.width = Inches(6.3)

    cell.text = text

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="D8E4BC"/>'.format(
            nsdecls('w')
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.font.name = "Candara"

            run.font.size = Pt(10)

# =========================================================
# ADD REPORT GENERAL INFO TABLE
# =========================================================

def add_report_general_info_table(
    document,
    report
):

    table = document.add_table(
        rows=6,
        cols=2
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    labels = [

        "Ruta del Reporte",

        "Modelo de Datos",

        "Frecuencia de Generación",

        "Tipo de Salida",

        "Archivos de Plantilla",

        "Origen de Datos"
    ]

    values = [

        report.get(
            "report_path",
            ""
        ),

        report.get(
            "data_model",
            ""
        ),

        report.get(
            "frequency",
            "No aplica"
        ),

        report.get(
            "output_format_string",
            ""
        ),

        report.get(
            "template_file_string",
            ""
        ),

        report.get(
            "datasource",
            ""
        )
    ]

    # =====================================================
    # ROWS
    # =====================================================

    for i in range(6):

        left = table.cell(i,0)
        right = table.cell(i,1)

        left.text = labels[i]
        right.text = values[i]

        apply_header_style(
            left,
            fill="D9D9D9"
        )


# =========================================================
# ADD PARAMETERS TABLE
# =========================================================

def add_parameters_table(
    document,
    parameters
):

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    headers = [

        "Nombre",

        "Código",

        "Obligatorio",

        "Descripción"
    ]

    hdr = table.rows[0].cells

    for i in range(4):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i],
            fill="808080",
            white_text=True
        )

    # =====================================================
    # PARAMETERS
    # =====================================================

    if not parameters:

        row = table.add_row().cells

        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""

        return

    for parameter in parameters:

        row = table.add_row().cells

        row[0].text = parameter.get(
            "name",
            ""
        )

        row[1].text = parameter.get(
            "code",
            ""
        )

        row[2].text = parameter.get(
            "mandatory",
            ""
        )

        row[3].text = parameter.get(
            "description",
            ""
        )


# =========================================================
# ADD REPORT OUTPUT SECTION
# =========================================================

def add_report_output_section(
    document
):

    p = document.add_paragraph()

    run = p.add_run(
        "•\tSalida del Reporte"
    )

    run.bold = True

    add_description_box(

        document,

        "Se presenta una imagen por cada "
        "salida según los tipos de salida "
        "indicados en el cuadro de Ficha "
        "general del reporte, para aquellos "
        "reportes que tengan una plantilla "
        "asociada."
    )

    document.add_paragraph("")


# =========================================================
# ADD SINGLE REPORT SECTION
# =========================================================

def add_single_report_section(
    document,
    report,
    section_number
):

    # =====================================================
    # TITLE
    # =====================================================

    create_header(

        document,

        f"5.{section_number}\t"
        f"{report.get('report_name','')}"
    )

    # =====================================================
    # GENERAL INFO
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        "•\tFicha General"
    )

    run.bold = True

    add_report_general_info_table(
        document,
        report
    )

    document.add_paragraph("")

    # =====================================================
    # PARAMETERS
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        "•\tParámetros de Ejecución"
    )

    run.bold = True

    add_parameters_table(

        document,

        report.get(
            "parameters",
            []
        )
    )

    document.add_paragraph("")

    # =====================================================
    # OUTPUT SECTION
    # =====================================================

    add_report_output_section(
        document
    )

    document.add_paragraph("")
    document.add_paragraph("")


# =========================================================
# ADD REPORT DESIGN SECTION
# =========================================================

def add_report_design_section(
    document,
    bip_metadata
):

    # =====================================================
    # SECTION TITLE
    # =====================================================

    create_header(
        document,
        "5\tDiseño del Reporte"
    )

    # =====================================================
    # DESCRIPTION
    # =====================================================

    add_description_box(

        document,

        "En esta sección se describe "
        "el detalle de cada uno de "
        "los reportes que conforman "
        "el desarrollo."
    )

    document.add_paragraph("")

    reports = bip_metadata.get(
        "reports",
        []
    )

    if not reports:

        document.add_paragraph(
            "No se encontraron reportes."
        )

        return

    # =====================================================
    # ITERATE REPORTS
    # =====================================================

    counter = 1

    for report in reports:

        add_single_report_section(

            document,

            report,

            counter
        )

        counter += 1