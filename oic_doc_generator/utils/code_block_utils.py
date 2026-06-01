# =========================================================
# FILE:
# oic_doc_generator/utils/code_block_utils.py
# =========================================================

from docx.shared import (
    Pt,
    RGBColor
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# =========================================================
# SQL KEYWORDS
# =========================================================

SQL_KEYWORDS = [

    "SELECT",
    "FROM",
    "WHERE",
    "AND",
    "OR",
    "ORDER",
    "GROUP",
    "BY",
    "INNER",
    "LEFT",
    "RIGHT",
    "JOIN",
    "ON",
    "UNION",
    "DISTINCT",
    "AS",
    "CASE",
    "WHEN",
    "THEN",
    "ELSE",
    "END",
    "EXISTS",
    "NOT",
    "NULL",
    "LIKE",
    "IN",
    "BETWEEN",
    "HAVING",
    "INSERT",
    "UPDATE",
    "DELETE",
    "CREATE",
    "ALTER",
    "DROP",
    "PACKAGE",
    "BODY",
    "PROCEDURE",
    "FUNCTION",
    "RETURN",
    "BEGIN",
    "EXCEPTION",
    "LOOP",
    "FOR",
    "WHILE",
    "IF"
]


# =========================================================
# CREATE SHADED PARAGRAPH
# =========================================================

def create_shaded_paragraph(
    document
):

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    shading = parse_xml(

        r'<w:shd {} w:fill="1E1E1E"/>'.format(
            nsdecls('w')
        )
    )

    paragraph._p.get_or_add_pPr().append(
        shading
    )

    return paragraph


# =========================================================
# ADD COLORED RUN
# =========================================================

def add_colored_run(
    paragraph,
    text,
    color
):

    run = paragraph.add_run(
        text
    )

    run.font.name = "Consolas"

    run.font.size = Pt(9)

    run.font.color.rgb = color

    return run


# =========================================================
# PROCESS TOKEN
# =========================================================

def process_token(
    paragraph,
    token
):

    clean = token.strip()

    upper = clean.upper()

    # =====================================================
    # KEYWORD
    # =====================================================

    if upper in SQL_KEYWORDS:

        add_colored_run(

            paragraph,

            token + " ",

            RGBColor(
                86,
                156,
                214
            )
        )

        return

    # =====================================================
    # VARIABLE
    # =====================================================

    if clean.startswith(":"):

        add_colored_run(

            paragraph,

            token + " ",

            RGBColor(
                78,
                201,
                176
            )
        )

        return

    # =====================================================
    # STRING
    # =====================================================

    if (
        clean.startswith("'")
        and
        clean.endswith("'")
    ):

        add_colored_run(

            paragraph,

            token + " ",

            RGBColor(
                214,
                157,
                133
            )
        )

        return

    # =====================================================
    # DEFAULT
    # =====================================================

    add_colored_run(

        paragraph,

        token + " ",

        RGBColor(
            220,
            220,
            220
        )
    )


# =========================================================
# ADD CODE BLOCK
# =========================================================

def add_code_block(
    document,
    text,
    language=None
):

    if not text:

        return

    paragraph = create_shaded_paragraph(
        document
    )

    lines = text.split(
        "\n"
    )

    for idx, line in enumerate(
        lines
    ):

        stripped = line.strip()

        # =================================================
        # COMMENTS
        # =================================================

        if stripped.startswith("--"):

            add_colored_run(

                paragraph,

                line,

                RGBColor(
                    106,
                    153,
                    85
                )
            )

        else:

            tokens = line.split(" ")

            for token in tokens:

                process_token(

                    paragraph,

                    token
                )

        # =================================================
        # NEXT LINE
        # =================================================

        if idx < len(lines) - 1:

            paragraph.add_run(
                "\n"
            )