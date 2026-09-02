# =========================================================
# FILE:
# oic_doc_generator/backend/utils/word_utils.py
# =========================================================

import re


from docx.shared import (
    Pt,
    RGBColor,
    Cm
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER
)

from docx.oxml import (
    parse_xml,
    OxmlElement
)

from docx.oxml.ns import (
    nsdecls,
    qn
)

from docx.enum.style import (
    WD_STYLE_TYPE
)


# =========================================================
# HEADER NUMBER REGEX
# =========================================================
#
# Ejemplos reconocidos:
#
# 1\tControl del Documento
# 1.1\tRegistro de Modificaciones
# 2\tVisión General
# 2.1\tReglas del Negocio
# 10.2\tTemas Cerrados
#
# =========================================================

HEADER_NUMBER_PATTERN = re.compile(

    r"^\s*"
    r"(\d+(?:\.\d+)*)"
    r"\t+"
    r"(.*)$"

)


# =========================================================
# GET NEXT NUMERIC ID
# =========================================================

def _get_next_numeric_id(
    elements,
    attribute_name
):

    values = []


    for element in elements:

        value = element.get(
            qn(
                attribute_name
            )
        )


        if value is None:

            continue


        try:

            values.append(
                int(
                    value
                )
            )

        except Exception:

            continue


    if not values:

        return 0


    return (
        max(
            values
        )
        + 1
    )


# =========================================================
# CREATE MULTILEVEL NUMBERING DEFINITION
# =========================================================

def _create_heading_abstract_numbering(
    document
):

    numbering = (
        document
        .part
        .numbering_part
        .element
    )


    abstract_num_id = (
        _get_next_numeric_id(

            numbering.findall(
                qn(
                    "w:abstractNum"
                )
            ),

            "w:abstractNumId"
        )
    )


    # =====================================================
    # ABSTRACT NUMBER
    # =====================================================

    abstract_num = OxmlElement(
        "w:abstractNum"
    )


    abstract_num.set(

        qn(
            "w:abstractNumId"
        ),

        str(
            abstract_num_id
        )
    )


    # =====================================================
    # MULTILEVEL
    # =====================================================

    multi_level_type = OxmlElement(
        "w:multiLevelType"
    )


    multi_level_type.set(

        qn(
            "w:val"
        ),

        "multilevel"
    )


    abstract_num.append(
        multi_level_type
    )


    # =====================================================
    # CREATE LEVELS
    # =====================================================
    #
    # Nivel 0:
    #   1
    #
    # Nivel 1:
    #   1.1
    #
    # Nivel 2:
    #   1.1.1
    #
    # etc.
    #
    # =====================================================

    for level in range(
        9
    ):

        lvl = OxmlElement(
            "w:lvl"
        )


        lvl.set(

            qn(
                "w:ilvl"
            ),

            str(
                level
            )
        )


        # ================================================
        # START
        # ================================================

        start = OxmlElement(
            "w:start"
        )


        start.set(

            qn(
                "w:val"
            ),

            "1"
        )


        lvl.append(
            start
        )


        # ================================================
        # NUMBER FORMAT
        # ================================================

        num_format = OxmlElement(
            "w:numFmt"
        )


        num_format.set(

            qn(
                "w:val"
            ),

            "decimal"
        )


        lvl.append(
            num_format
        )


        # ================================================
        # NUMBER FORMAT TEXT
        # ================================================
        #
        # 0 -> %1
        # 1 -> %1.%2
        # 2 -> %1.%2.%3
        #
        # ================================================

        level_text_value = ".".join(

            f"%{index}"

            for index in range(
                1,
                level + 2
            )
        )


        level_text = OxmlElement(
            "w:lvlText"
        )


        level_text.set(

            qn(
                "w:val"
            ),

            level_text_value
        )


        lvl.append(
            level_text
        )


        # ================================================
        # TAB AFTER NUMBER
        # ================================================

        suffix = OxmlElement(
            "w:suff"
        )


        suffix.set(

            qn(
                "w:val"
            ),

            "tab"
        )


        lvl.append(
            suffix
        )


        # ================================================
        # ALIGN NUMBER LEFT
        # ================================================

        level_justification = OxmlElement(
            "w:lvlJc"
        )


        level_justification.set(

            qn(
                "w:val"
            ),

            "left"
        )


        lvl.append(
            level_justification
        )


        # ================================================
        # PARAGRAPH POSITION
        # ================================================
        #
        # 720 twips = 0.5 pulgadas.
        #
        # Esto reproduce aproximadamente la tabulación que
        # actualmente existe entre:
        #
        # 2     Visión General
        #
        # ================================================

        paragraph_properties = OxmlElement(
            "w:pPr"
        )


        tabs = OxmlElement(
            "w:tabs"
        )


        tab = OxmlElement(
            "w:tab"
        )


        tab.set(

            qn(
                "w:val"
            ),

            "num"
        )


        tab.set(

            qn(
                "w:pos"
            ),

            "720"
        )


        tabs.append(
            tab
        )


        paragraph_properties.append(
            tabs
        )


        indent = OxmlElement(
            "w:ind"
        )


        indent.set(

            qn(
                "w:left"
            ),

            "720"
        )


        indent.set(

            qn(
                "w:hanging"
            ),

            "720"
        )


        paragraph_properties.append(
            indent
        )


        lvl.append(
            paragraph_properties
        )


        abstract_num.append(
            lvl
        )


    numbering.append(
        abstract_num
    )


    return abstract_num_id


# =========================================================
# CREATE NUMBERING INSTANCE
# =========================================================

def _create_heading_numbering_instance(
    document,
    abstract_num_id,
    start_at=1
):

    numbering = (
        document
        .part
        .numbering_part
        .element
    )


    num_id = (
        _get_next_numeric_id(

            numbering.findall(
                qn(
                    "w:num"
                )
            ),

            "w:numId"
        )
    )


    # =====================================================
    # NUM
    # =====================================================

    num = OxmlElement(
        "w:num"
    )


    num.set(

        qn(
            "w:numId"
        ),

        str(
            num_id
        )
    )


    # =====================================================
    # ABSTRACT REFERENCE
    # =====================================================

    abstract_reference = OxmlElement(
        "w:abstractNumId"
    )


    abstract_reference.set(

        qn(
            "w:val"
        ),

        str(
            abstract_num_id
        )
    )


    num.append(
        abstract_reference
    )


    # =====================================================
    # OPTIONAL START OVERRIDE
    # =====================================================
    #
    # Esto permite conservar casos como:
    #
    # 8 Control de Accesos
    # 10 Temas Abiertos y Cerrados
    #
    # Si deliberadamente no existe una sección 9,
    # Word no cambiará el 10 por un 9.
    #
    # =====================================================

    if start_at != 1:

        level_override = OxmlElement(
            "w:lvlOverride"
        )


        level_override.set(

            qn(
                "w:ilvl"
            ),

            "0"
        )


        start_override = OxmlElement(
            "w:startOverride"
        )


        start_override.set(

            qn(
                "w:val"
            ),

            str(
                start_at
            )
        )


        level_override.append(
            start_override
        )


        num.append(
            level_override
        )


    numbering.append(
        num
    )


    return num_id


# =========================================================
# GET DS140 NUMBERING STATE
# =========================================================

def _get_heading_numbering_state(
    document
):

    state = getattr(

        document,

        "_ds140_heading_numbering",

        None
    )


    if state is not None:

        return state


    abstract_num_id = (
        _create_heading_abstract_numbering(
            document
        )
    )


    num_id = (
        _create_heading_numbering_instance(

            document,

            abstract_num_id,

            start_at=1
        )
    )


    state = {

        "abstract_num_id":
            abstract_num_id,

        "num_id":
            num_id,

        "last_top_level":
            0
    }


    setattr(

        document,

        "_ds140_heading_numbering",

        state
    )


    return state


# =========================================================
# APPLY NUMBERING TO PARAGRAPH
# =========================================================

def _apply_heading_numbering(
    paragraph,
    num_id,
    level,
    size
):

    paragraph_properties = (
        paragraph
        ._p
        .get_or_add_pPr()
    )


    # =====================================================
    # NUMBER PROPERTIES
    # =====================================================

    num_properties = (
        paragraph_properties
        .get_or_add_numPr()
    )


    level_element = (
        num_properties
        .get_or_add_ilvl()
    )


    level_element.val = (
        level - 1
    )


    num_id_element = (
        num_properties
        .get_or_add_numId()
    )


    num_id_element.val = (
        num_id
    )


    # =====================================================
    # PARAGRAPH MARK FORMAT
    # =====================================================
    #
    # La numeración automática de Word utiliza el formato
    # del paragraph mark.
    #
    # Esto hace que el número tenga el mismo Arial,
    # negrita y tamaño que el título.
    #
    # =====================================================

    paragraph_run_properties = (
        paragraph_properties.find(
            qn(
                "w:rPr"
            )
        )
    )


    if paragraph_run_properties is None:

        paragraph_run_properties = (
            OxmlElement(
                "w:rPr"
            )
        )


        paragraph_properties.append(
            paragraph_run_properties
        )


    fonts = OxmlElement(
        "w:rFonts"
    )


    fonts.set(

        qn(
            "w:ascii"
        ),

        "Arial"
    )


    fonts.set(

        qn(
            "w:hAnsi"
        ),

        "Arial"
    )


    fonts.set(

        qn(
            "w:eastAsia"
        ),

        "Arial"
    )


    paragraph_run_properties.append(
        fonts
    )


    bold = OxmlElement(
        "w:b"
    )


    paragraph_run_properties.append(
        bold
    )


    font_size = OxmlElement(
        "w:sz"
    )


    font_size.set(

        qn(
            "w:val"
        ),

        str(
            int(
                size * 2
            )
        )
    )


    paragraph_run_properties.append(
        font_size
    )


    font_size_complex = OxmlElement(
        "w:szCs"
    )


    font_size_complex.set(

        qn(
            "w:val"
        ),

        str(
            int(
                size * 2
            )
        )
    )


    paragraph_run_properties.append(
        font_size_complex
    )


# =========================================================
# GET TABLE OF CONTENTS STATE
# =========================================================

def _get_toc_state(
    document
):

    state = getattr(
        document,
        "_ds140_toc_state",
        None
    )


    if state is not None:

        return state


    state = {

        "entries": [],

        "bookmark_counter": 1
    }


    setattr(
        document,
        "_ds140_toc_state",
        state
    )


    return state


# =========================================================
# REGISTER HEADING FOR TABLE OF CONTENTS
# =========================================================

def _register_toc_heading(
    document,
    paragraph,
    run,
    number_text,
    title,
    level
):

    state = _get_toc_state(
        document
    )


    bookmark_id = (
        state[
            "bookmark_counter"
        ]
    )


    state[
        "bookmark_counter"
    ] += 1


    bookmark_name = (
        f"ds140_heading_{bookmark_id}"
    )


    # =====================================================
    # BOOKMARK START
    # =====================================================

    bookmark_start = OxmlElement(
        "w:bookmarkStart"
    )


    bookmark_start.set(
        qn("w:id"),
        str(
            bookmark_id
        )
    )


    bookmark_start.set(
        qn("w:name"),
        bookmark_name
    )


    # =====================================================
    # BOOKMARK END
    # =====================================================

    bookmark_end = OxmlElement(
        "w:bookmarkEnd"
    )


    bookmark_end.set(
        qn("w:id"),
        str(
            bookmark_id
        )
    )


    # =====================================================
    # PLACE BOOKMARK AROUND TITLE
    # =====================================================

    run._r.addprevious(
        bookmark_start
    )


    run._r.addnext(
        bookmark_end
    )


    # =====================================================
    # REGISTER ENTRY
    # =====================================================

    state[
        "entries"
    ].append({

        "number":
            number_text,

        "title":
            title,

        "text":
            f"{number_text} {title}",

        "level":
            level,

        "bookmark":
            bookmark_name
    })


# =========================================================
# CREATE INTERNAL HYPERLINK
# =========================================================

def _add_internal_hyperlink(
    paragraph,
    text,
    bookmark_name
):

    hyperlink = OxmlElement(
        "w:hyperlink"
    )

    hyperlink.set(
        qn("w:anchor"),
        bookmark_name
    )

    hyperlink.set(
        qn("w:history"),
        "1"
    )


    run = OxmlElement(
        "w:r"
    )

    run_properties = OxmlElement(
        "w:rPr"
    )


    # =====================================================
    # FONT
    # =====================================================

    fonts = OxmlElement(
        "w:rFonts"
    )

    fonts.set(
        qn("w:ascii"),
        "Arial"
    )

    fonts.set(
        qn("w:hAnsi"),
        "Arial"
    )

    fonts.set(
        qn("w:eastAsia"),
        "Arial"
    )

    run_properties.append(
        fonts
    )


    # =====================================================
    # FONT SIZE
    # 9 PT = 18 half-points
    # =====================================================

    size = OxmlElement(
        "w:sz"
    )

    size.set(
        qn("w:val"),
        "18"
    )

    run_properties.append(
        size
    )


    size_cs = OxmlElement(
        "w:szCs"
    )

    size_cs.set(
        qn("w:val"),
        "18"
    )

    run_properties.append(
        size_cs
    )


    # =====================================================
    # BLACK
    # =====================================================

    color = OxmlElement(
        "w:color"
    )

    color.set(
        qn("w:val"),
        "000000"
    )

    run_properties.append(
        color
    )


    # =====================================================
    # NO UNDERLINE
    # =====================================================

    underline = OxmlElement(
        "w:u"
    )

    underline.set(
        qn("w:val"),
        "none"
    )

    run_properties.append(
        underline
    )


    run.append(
        run_properties
    )


    text_element = OxmlElement(
        "w:t"
    )

    text_element.text = (
        text
    )

    run.append(
        text_element
    )


    hyperlink.append(
        run
    )

    paragraph._p.append(
        hyperlink
    )


# =========================================================
# ADD PAGE REFERENCE FIELD
# =========================================================

def _add_pageref_field(
    paragraph,
    bookmark_name
):

    run = paragraph.add_run()


    # =====================================================
    # FIELD BEGIN
    # =====================================================

    field_begin = OxmlElement(
        "w:fldChar"
    )

    field_begin.set(
        qn("w:fldCharType"),
        "begin"
    )

    field_begin.set(
        qn("w:dirty"),
        "true"
    )


    # =====================================================
    # FIELD INSTRUCTION
    # =====================================================

    instruction = OxmlElement(
        "w:instrText"
    )

    instruction.set(
        qn("xml:space"),
        "preserve"
    )

    instruction.text = (
        f" PAGEREF {bookmark_name} \\h "
    )


    # =====================================================
    # FIELD SEPARATOR
    # =====================================================

    field_separator = OxmlElement(
        "w:fldChar"
    )

    field_separator.set(
        qn("w:fldCharType"),
        "separate"
    )


    # =====================================================
    # INITIAL RESULT
    # =====================================================

    result_run = OxmlElement(
        "w:r"
    )

    result_properties = OxmlElement(
        "w:rPr"
    )


    fonts = OxmlElement(
        "w:rFonts"
    )

    fonts.set(
        qn("w:ascii"),
        "Arial"
    )

    fonts.set(
        qn("w:hAnsi"),
        "Arial"
    )

    result_properties.append(
        fonts
    )


    size = OxmlElement(
        "w:sz"
    )

    size.set(
        qn("w:val"),
        "18"
    )

    result_properties.append(
        size
    )


    result_run.append(
        result_properties
    )


    result_text = OxmlElement(
        "w:t"
    )

    # Word reemplazará este valor al actualizar el campo.
    result_text.text = "1"

    result_run.append(
        result_text
    )


    # =====================================================
    # FIELD END
    # =====================================================

    field_end = OxmlElement(
        "w:fldChar"
    )

    field_end.set(
        qn("w:fldCharType"),
        "end"
    )


    run._r.append(
        field_begin
    )

    run._r.append(
        instruction
    )

    run._r.append(
        field_separator
    )

    paragraph._p.append(
        result_run
    )

    end_run = paragraph.add_run()

    end_run._r.append(
        field_end
    )


# =========================================================
# CREATE TABLE OF CONTENTS PLACEHOLDER
# =========================================================

def create_toc_table(
    document
):

    # =====================================================
    # CONTENT LABEL
    # =====================================================

    content_title = document.add_paragraph()

    content_title.paragraph_format.space_before = Pt(
        18
    )

    content_title.paragraph_format.space_after = Pt(
        4
    )


    run = content_title.add_run(
        "Contenido"
    )

    run.bold = True

    run.font.name = (
        "Arial"
    )

    run.font.size = Pt(
        11
    )

    run.font.color.rgb = RGBColor(
        47,
        85,
        151
    )


    # =====================================================
    # PLACEHOLDER
    # =====================================================

    placeholder = document.add_paragraph()

    return placeholder


# =========================================================
# POPULATE TABLE OF CONTENTS
# =========================================================

def populate_toc_table(
    document,
    table
):

    # "table" ahora realmente es el paragraph placeholder.
    placeholder = table


    state = _get_toc_state(
        document
    )


    entries = state.get(
        "entries",
        []
    )


    if not entries:

        paragraph = (
            placeholder
            .insert_paragraph_before()
        )

        run = paragraph.add_run(
            "No se encontraron secciones."
        )

        run.font.name = (
            "Arial"
        )

        run.font.size = Pt(
            9
        )

        return


    # =====================================================
    # AVAILABLE PAGE WIDTH
    # =====================================================

    section = document.sections[
        0
    ]


    available_width = (

        section.page_width

        -

        section.left_margin

        -

        section.right_margin
    )


    # Leave a little space before right margin.
    page_number_position = (

        available_width

        -

        Cm(
            0.15
        )
    )


    # =====================================================
    # CREATE ENTRIES
    # =====================================================

    for entry in entries:

        number_text = entry.get(
            "number",
            ""
        )

        title = entry.get(
            "title",
            ""
        )

        level = entry.get(
            "level",
            1
        )

        bookmark = entry.get(
            "bookmark",
            ""
        )


        paragraph = (
            placeholder
            .insert_paragraph_before()
        )


        # =================================================
        # BASIC FORMAT
        # =================================================

        paragraph.paragraph_format.space_before = Pt(
            0
        )

        paragraph.paragraph_format.space_after = Pt(
            2
        )

        paragraph.paragraph_format.line_spacing = (
            1
        )


        # =================================================
        # LEVEL POSITION
        # =================================================

        if level == 1:

            paragraph.paragraph_format.left_indent = Cm(
                0
            )

            title_position = Cm(
                0.85
            )

        else:

            paragraph.paragraph_format.left_indent = Cm(
                0.35
            )

            title_position = Cm(
                1.65
            )


        # =================================================
        # TAB STOPS
        # =================================================

        tab_stops = (
            paragraph
            .paragraph_format
            .tab_stops
        )


        # Position where title begins after section number.
        if number_text:

            tab_stops.add_tab_stop(

                title_position,

                WD_TAB_ALIGNMENT.LEFT
            )


        # Right tab with dot leader.
        tab_stops.add_tab_stop(

            page_number_position,

            WD_TAB_ALIGNMENT.RIGHT,

            WD_TAB_LEADER.DOTS
        )


        # =================================================
        # NUMBER
        # =================================================

        if number_text:

            number_run = paragraph.add_run(
                number_text
            )

            number_run.font.name = (
                "Arial"
            )

            number_run.font.size = Pt(
                9
            )


            number_run.add_tab()


        # =================================================
        # TITLE / INTERNAL LINK
        # =================================================

        _add_internal_hyperlink(

            paragraph,

            title,

            bookmark
        )


        # =================================================
        # DOT LEADER
        # =================================================

        separator_run = (
            paragraph.add_run()
        )

        separator_run.add_tab()


        # =================================================
        # PAGE NUMBER
        # =================================================

        _add_pageref_field(

            paragraph,

            bookmark
        )


    # =====================================================
    # REMOVE PLACEHOLDER
    # =====================================================

    placeholder_element = (
        placeholder._element
    )

    placeholder_element.getparent().remove(
        placeholder_element
    )

# =========================================================
# CREATE HEADER
# =========================================================

def create_header(
    document,
    text,
    size=16
):

    # =====================================================
    # DETECT REAL NUMBER
    # =====================================================

    match = HEADER_NUMBER_PATTERN.match(
        text
    )


    # =====================================================
    # NUMBERED HEADER
    # =====================================================

    if match:

        number_text = (
            match
            .group(
                1
            )
        )


        title = (
            match
            .group(
                2
            )
        )


        numbers = [

            int(
                value
            )

            for value in number_text.split(
                "."
            )
        ]


        level = len(
            numbers
        )


        # =================================================
        # STYLE
        # =================================================

        if level == 1:

            p = document.add_paragraph(
                style="Heading 1"
            )


            p.style = (
                "HD1"
            )

        else:

            p = document.add_paragraph(
                style="Heading 2"
            )


            p.style = (
                "HD2"
            )


        # =================================================
        # NUMBERING STATE
        # =================================================

        state = (
            _get_heading_numbering_state(
                document
            )
        )


        # =================================================
        # TOP LEVEL
        # =================================================
        #
        # Normalmente Word seguirá:
        #
        # 1
        # 2
        # 3
        # ...
        #
        # Pero si detectamos un salto deliberado:
        #
        # 8
        # 10
        #
        # se crea una nueva instancia que comienza en 10.
        #
        # =================================================

        if level == 1:

            explicit_number = (
                numbers[
                    0
                ]
            )


            if (
                state[
                    "last_top_level"
                ]
                ==
                0
            ):

                expected_number = 1

            else:

                expected_number = (

                    state[
                        "last_top_level"
                    ]
                    + 1
                )


            if (
                explicit_number
                !=
                expected_number
            ):

                state[
                    "num_id"
                ] = (
                    _create_heading_numbering_instance(

                        document,

                        state[
                            "abstract_num_id"
                        ],

                        start_at=
                            explicit_number
                    )
                )


            state[
                "last_top_level"
            ] = (
                explicit_number
            )


        # =================================================
        # APPLY REAL NUMBERING
        # =================================================

        _apply_heading_numbering(

            p,

            state[
                "num_id"
            ],

            level,

            size
        )


        # =================================================
        # ADD TITLE ONLY
        # =================================================
        #
        # IMPORTANTE:
        #
        # Ya NO agregamos:
        #
        # 2.1\tReglas del Negocio
        #
        # Solo agregamos:
        #
        # Reglas del Negocio
        #
        # El 2.1 será generado por Word.
        #
        # =================================================

        run = p.add_run(
            title
        )

        # =================================================
        # REGISTER IN TABLE OF CONTENTS
        # =================================================

        _register_toc_heading(

            document,

            p,

            run,

            number_text,

            title,

            level
        )


    # =====================================================
    # NON NUMBERED HEADER
    # =====================================================

    else:

        if "." in text:

            p = document.add_paragraph(
                style="Heading 2"
            )


            p.style = (
                "HD2"
            )

        else:

            p = document.add_paragraph(
                style="Heading 1"
            )


            p.style = (
                "HD1"
            )


        run = p.add_run(
            text
        )

        # =================================================
        # TABLE OF CONTENTS TITLE
        # =================================================

        if text == "Tabla de Contenido":

            _register_toc_heading(

                document,

                p,

                run,

                "",

                text,

                1
            )


    # =====================================================
    # CURRENT HEADER FORMAT
    # =====================================================

    run.bold = True

    run.font.name = (
        "Arial"
    )

    run.font.size = Pt(
        size
    )


    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )


    return p


# =========================================================
# ADD DESCRIPTION BOX
# =========================================================

def add_description_box(
    document,
    text
):

    table = document.add_table(
        rows=1,
        cols=1
    )


    table.autofit = False


    cell = (
        table
        .rows[
            0
        ]
        .cells[
            0
        ]
    )


    cell.width = Pt(
        450
    )


    shading = parse_xml(

        r'<w:shd {} w:fill="D8E4BC"/>'.format(
            nsdecls(
                'w'
            )
        )
    )


    cell._tc.get_or_add_tcPr().append(
        shading
    )


    p = cell.paragraphs[
        0
    ]


    run = p.add_run(
        text
    )


    run.font.name = (
        "Candara"
    )


    run.font.size = Pt(
        10
    )


    run.font.color.rgb = RGBColor(
        0,
        0,
        0
    )


    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )

# =========================================================
# CREATE DOCUMENT STYLES
# =========================================================

def create_document_styles(
    document
):

    styles = document.styles


    # =====================================================
    # HD1
    # =====================================================

    if "HD1" not in styles:

        style = styles.add_style(
            "HD1",
            WD_STYLE_TYPE.PARAGRAPH
        )

        font = style.font

        font.name = "Arial"
        font.size = Pt(16)
        font.bold = True

        paragraph_format = (
            style.paragraph_format
        )

        paragraph_format.left_indent = Cm(0)
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(12)
        paragraph_format.keep_with_next = True
        paragraph_format.keep_together = True

        p_pr = (
            style.element
            .get_or_add_pPr()
        )

        outline_level = OxmlElement(
            "w:outlineLvl"
        )

        outline_level.set(
            qn("w:val"),
            "0"
        )

        p_pr.append(
            outline_level
        )

        border_xml = parse_xml(r'''
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top
                    w:val="single"
                    w:sz="45"
                    w:space="1"
                    w:color="000000"/>
            </w:pBdr>
        ''')

        p_pr.append(
            border_xml
        )


    # =====================================================
    # HD2
    # =====================================================

    if "HD2" not in styles:

        style = styles.add_style(
            "HD2",
            WD_STYLE_TYPE.PARAGRAPH
        )

        font = style.font

        font.name = "Arial"
        font.size = Pt(13)
        font.bold = True

        paragraph_format = (
            style.paragraph_format
        )

        paragraph_format.left_indent = Cm(0)
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.keep_with_next = True
        paragraph_format.keep_together = True

        p_pr = (
            style.element
            .get_or_add_pPr()
        )

        outline_level = OxmlElement(
            "w:outlineLvl"
        )

        outline_level.set(
            qn("w:val"),
            "1"
        )

        p_pr.append(
            outline_level
        )

        border_xml = parse_xml(r'''
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top
                    w:val="single"
                    w:sz="24"
                    w:space="1"
                    w:color="000000"/>
            </w:pBdr>
        ''')

        p_pr.append(
            border_xml
        )


# =========================================================
# APPLY TABLE HEADER STYLE
# =========================================================

def apply_table_header_style(
    cell,
    fill="D9D9D9"
):

    shading = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(
            nsdecls("w"),
            fill
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading
    )

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)