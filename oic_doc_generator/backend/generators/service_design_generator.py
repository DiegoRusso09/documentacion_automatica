# =========================================================
# FILE: service_design_generator.py
# =========================================================

from docx.shared import Pt

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.enum.table import (
    WD_TABLE_ALIGNMENT
)

from docx.oxml import parse_xml

from docx.oxml.ns import (
    nsdecls
)

from oic_doc_generator.backend.parsers.service_parser import (
    get_request_response_schemas
)


# =========================================================
# HEADER STYLE
# =========================================================

def style_header_cell(cell):

    for paragraph in cell.paragraphs:

        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        if paragraph.runs:

            paragraph.runs[0].bold = True

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="D9D9D9"/>'.format(
            nsdecls('w')
        )

    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )


# =========================================================
# BUILD JSON PAYLOAD
# =========================================================

def build_json_payload(
    elements
):

    if not elements:

        return "No aplica."

    lines = []

    lines.append("{")

    for element in elements:

        name = element.get(
            "name",
            "field"
        )

        data_type = element.get(
            "type",
            "string"
        )

        lines.append(
            f'  "{name}": "{data_type}",'
        )

    # =====================================================
    # REMOVE LAST COMMA
    # =====================================================

    if len(lines) > 1:

        lines[-1] = lines[-1].rstrip(",")

    lines.append("}")

    return "\n".join(lines)

# =========================================================
# CREATE SERVICE TABLE
# =========================================================

def create_service_table(
    document,
    request_elements,
    response_elements
):

    table = document.add_table(
        rows=4,
        cols=1
    )

    table.style = "Table Grid"

    # =====================================================
    # REQUEST HEADER
    # =====================================================

    request_header = table.cell(
        0,
        0
    )

    request_header.text = "REQUEST"

    style_header_cell(
        request_header
    )

    # =====================================================
    # REQUEST BODY
    # =====================================================

    request_body = table.cell(
        1,
        0
    )

    request_body.text = build_json_payload(
        request_elements
    )

    # =====================================================
    # RESPONSE HEADER
    # =====================================================

    response_header = table.cell(
        2,
        0
    )

    response_header.text = "RESPONSE"

    style_header_cell(
        response_header
    )

    # =====================================================
    # RESPONSE BODY
    # =====================================================

    response_body = table.cell(
        3,
        0
    )

    response_body.text = build_json_payload(
        response_elements
    )
    
# =========================================================
# ADD SERVICE DESIGN SECTION
# =========================================================

def add_service_design_section(
    document,
    integration_path,
    integration_name,
    version
):

    services = get_request_response_schemas(
        integration_path
    )

    # =====================================================
    # FILTER VALID SERVICES
    # =====================================================

    valid_services = []

    for service in services:

        request_elements = service.get(
            "request_elements",
            []
        )

        response_elements = service.get(
            "response_elements",
            []
        )

        if (

            not request_elements

            and

            not response_elements
        ):

            continue

        valid_services.append(
            service
        )

    # =====================================================
    # NO VALID SERVICES
    # =====================================================

    if not valid_services:

        return

    # =====================================================
    # INTEGRATION TITLE
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        f"• {integration_name}|{version}"
    )

    run.bold = True

    # =====================================================
    # ITERATE SERVICES
    # =====================================================

    for service in valid_services:

        endpoint_name = service.get(
            "endpoint",
            "UNKNOWN"
        )

        request_elements = service.get(
            "request_elements",
            []
        )

        response_elements = service.get(
            "response_elements",
            []
        )

        # =================================================
        # ENDPOINT TITLE
        # =================================================

        endpoint_paragraph = (
            document.add_paragraph()
        )

        endpoint_run = (
            endpoint_paragraph.add_run(
                endpoint_name
            )
        )

        endpoint_run.bold = True

        # =================================================
        # DATA TYPE
        # =================================================

        document.add_paragraph(
            "Tipo de Dato\t: JSON"
        )

        # =================================================
        # TABLE
        # =================================================

        create_service_table(

            document,

            request_elements,

            response_elements
        )

        document.add_paragraph("")