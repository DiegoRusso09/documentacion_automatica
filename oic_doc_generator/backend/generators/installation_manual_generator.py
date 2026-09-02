# =========================================================
# FILE:
# oic_doc_generator/backend/generators/
# installation_manual_generator.py
# =========================================================

from io import BytesIO
from datetime import datetime
import os

from docx import Document

from docx.shared import (
    Cm,
    Pt,
    RGBColor
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.enum.table import (
    WD_TABLE_ALIGNMENT
)

from docx.oxml import (
    parse_xml,
    OxmlElement
)

from docx.oxml.ns import (
    nsdecls,
    qn
)

from oic_doc_generator.backend.utils.word_utils import (
    create_document_styles,
    apply_table_header_style,
    create_header,
    create_toc_table,
    populate_toc_table
)

from oic_doc_generator.backend.renderers.bip_installation_renderer import (
    render_bip_path_image,
    render_bip_tasks_image,
    render_bip_upload_image,
    render_bip_validation_image
)

# =========================================================
# TEMPLATE PATH
# =========================================================

def get_im090_template_path(
    file_name
):

    current_dir = os.path.dirname(
        os.path.abspath(
            __file__
        )
    )


    project_root = os.path.dirname(
        os.path.dirname(
            current_dir
        )
    )


    return os.path.join(
        project_root,
        "templates",
        file_name
    )


# =========================================================
# ADD CENTERED IMAGE
# =========================================================

def add_centered_image(
    document,
    image_path,
    width
):

    if (
        not image_path
        or
        not os.path.exists(
            image_path
        )
    ):

        return


    paragraph = (
        document.add_paragraph()
    )


    paragraph.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )


    run = (
        paragraph.add_run()
    )


    run.add_picture(
        image_path,
        width=width
    )


# =========================================================
# BIP INSTALLATION PLAN TABLE
# =========================================================

def add_bip_installation_plan_table(
    document,
    items
):

    if not items:

        return


    table = document.add_table(
        rows=1,
        cols=4
    )


    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )


    headers = [
        "Orden",
        "Tipo",
        "Archivo",
        "Ruta de Instalación"
    ]


    for index, header in enumerate(
        headers
    ):

        cell = table.cell(
            0,
            index
        )

        cell.text = header

        apply_table_header_style(
            cell
        )


    for item in items:

        cells = (
            table.add_row().cells
        )


        cells[0].text = str(
            item.get(
                "order",
                ""
            )
        )


        cells[1].text = (
            item.get(
                "type_label",
                ""
            )
        )


        cells[2].text = (
            item.get(
                "file_name",
                ""
            )
        )


        cells[3].text = (
            item.get(
                "display_upload_path",
                ""
            )
        )


        for cell in cells:

            for paragraph in cell.paragraphs:

                for run in paragraph.runs:

                    run.font.name = "Arial"

                    run.font.size = Pt(
                        8
                    )


    for row in table.rows:

        row.cells[0].width = Cm(
            1.2
        )

        row.cells[1].width = Cm(
            2.7
        )

        row.cells[2].width = Cm(
            4.2
        )

        row.cells[3].width = Cm(
            7.0
        )


    return table


# =========================================================
# BI PUBLISHER INSTALLATION SECTION
# =========================================================

def add_bip_installation_section(
    document,
    bip_installation_plan
):

    if not bip_installation_plan:

        return


    items = (
        bip_installation_plan.get(
            "items",
            []
        )
    )


    if not items:

        return


    # =====================================================
    # 3.2
    # =====================================================

    create_header(
        document,
        "3.2\tERP Cloud - Configuración de BI Publisher",
        size=14
    )


    # =====================================================
    # 3.2.1 ACCESS
    # =====================================================

    create_header(
        document,
        "3.2.1\tAcceso al Catálogo de BI Publisher",
        size=12
    )


    document.add_paragraph(
        (
            "Acceda al ambiente destino de Oracle ERP Cloud "
            "con un usuario que cuente con los privilegios "
            "requeridos para administrar los artefactos "
            "de BI Publisher."
        )
    )


    p = document.add_paragraph(
        "Ingrese al Catálogo de BI Publisher mediante el siguiente enlace:"
    )


    p = document.add_paragraph()


    run = p.add_run(
        "<ERP Domain>"
    )

    run.font.color.rgb = RGBColor(
        255,
        0,
        0
    )


    run = p.add_run(
        "/xmlpserver/servlet/catalog"
    )

    run.font.color.rgb = RGBColor(
        0,
        0,
        255
    )


    signin_path = (
        get_im090_template_path(
            "singin.png"
        )
    )


    add_centered_image(
        document,
        signin_path,
        Cm(12.5)
    )


    # =====================================================
    # 3.2.2 CONSIDERATIONS
    # =====================================================

    create_header(
        document,
        "3.2.2\tConsideraciones para la instalación",
        size=12
    )


    considerations = [

        (
            "Los artefactos deberán instalarse respetando "
            "el orden establecido en este documento."
        ),

        (
            "Las carpetas se instalarán antes que los "
            "Data Models y los Reportes."
        ),

        (
            "La ruta indicada para cada artefacto "
            "corresponde al Catálogo de BI Publisher "
            "del ambiente destino."
        ),

        (
            "No modifique el nombre ni el contenido de "
            "los archivos incluidos en la carpeta OTBI "
            "del paquete de instalación."
        ),

        (
            "Si una ruta requerida no existe, deberá "
            "crearse antes de continuar, excepto cuando "
            "dicha carpeta sea provista mediante un "
            "artefacto XDRZ incluido en el mismo pase."
        )
    ]


    for consideration in considerations:

        document.add_paragraph(
            consideration,
            style="List Bullet"
        )


        document.add_paragraph("")


        p = document.add_paragraph()

        run = p.add_run(
            "Orden de instalación"
        )

        run.bold = True


        document.add_paragraph(
            (
                "Los artefactos incluidos en la carpeta OTBI "
                "deberán instalarse respetando el siguiente orden:"
            )
        )


        add_bip_installation_plan_table(
            document,
            items
        )


        document.add_paragraph("")

    # =====================================================
    # ARTIFACTS
    # =====================================================

    subsection = 3


    for item in items:

        artifact_type = (
            item.get(
                "artifact_type",
                ""
            )
        )


        type_label = (
            item.get(
                "type_label",
                "Artefacto"
            )
        )


        object_name = (
            item.get(
                "object_name",
                ""
            )
        )


        file_name = (
            item.get(
                "file_name",
                ""
            )
        )


        package_path = (
            item.get(
                "package_path",
                ""
            )
        )


        display_upload_path = (
            item.get(
                "display_upload_path",
                ""
            )
        )


        target_folder_name = (
            item.get(
                "target_folder_name",
                ""
            )
        )


        # =================================================
        # HEADER
        # =================================================

        create_header(

            document,

            (
                f"3.2.{subsection}\t"
                f"Instalación de {type_label}: "
                f"{object_name}"
            ),

            size=12
        )


        subsection += 1


        # =================================================
        # STEP 1 - ROUTE
        # =================================================

        p = document.add_paragraph()

        run = p.add_run(
            "1. Ubicación de la ruta de destino"
        )

        run.bold = True


        document.add_paragraph(
            (
                "En el Catálogo de BI Publisher, "
                "navegue hasta la siguiente ruta:"
            )
        )


        p = document.add_paragraph()

        run = p.add_run(
            display_upload_path
        )

        run.font.color.rgb = RGBColor(
            0,
            0,
            255
        )


        route_image = (
            render_bip_path_image(
                display_upload_path
            )
        )


        add_centered_image(
            document,
            route_image,
            Cm(15)
        )


        # =================================================
        # STEP 2 - TASKS
        # =================================================

        p = document.add_paragraph()

        run = p.add_run(
            "2. Selección de la opción Cargar"
        )

        run.bold = True


        document.add_paragraph(
            (
                f'Una vez ubicado en la carpeta '
                f'"{target_folder_name}", seleccione '
                f'la opción "Cargar" disponible '
                f'en el menú de tareas.'
            )
        )


        tasks_image = (
            render_bip_tasks_image(
                target_folder_name
            )
        )


        add_centered_image(
            document,
            tasks_image,
            Cm(10.5)
        )


        # =================================================
        # STEP 3 - FILE
        # =================================================

        p = document.add_paragraph()

        run = p.add_run(
            "3. Selección del archivo de instalación"
        )

        run.bold = True


        document.add_paragraph(
            (
                "En la ventana de carga, seleccione "
                "el siguiente archivo incluido en el "
                "paquete de instalación:"
            )
        )


        p = document.add_paragraph()

        run = p.add_run(
            package_path
        )

        run.bold = True


        upload_image = (
            render_bip_upload_image(
                file_name
            )
        )


        add_centered_image(
            document,
            upload_image,
            Cm(15)
        )


        # =================================================
        # STEP 4 - UPLOAD
        # =================================================

        p = document.add_paragraph()

        run = p.add_run(
            "4. Ejecución de la carga"
        )

        run.bold = True


        document.add_paragraph(
            (
                f'Seleccione el archivo "{file_name}" '
                f'y haga clic en "Cargar". '
                f'Espere hasta que la operación '
                f'finalice antes de continuar con '
                f'el siguiente artefacto.'
            )
        )


        # =================================================
        # STEP 5 - VALIDATION
        # =================================================

        p = document.add_paragraph()

        run = p.add_run(
            "5. Validación de la instalación"
        )

        run.bold = True


        # =================================================
        # VALIDATION TEXT
        # =================================================

        if artifact_type == "folder":

            validation_text = (

                f'Verifique que la carpeta '
                f'"{object_name}" haya sido creada '
                f'correctamente en el Catálogo de '
                f'BI Publisher.'
            )


        elif artifact_type == "datamodel":

            validation_text = (

                f'Verifique que el Data Model '
                f'"{object_name}" se encuentre '
                f'disponible en la ruta indicada.'
            )


        else:

            validation_text = (

                f'Verifique que el Reporte '
                f'"{object_name}" se encuentre '
                f'disponible en la ruta indicada.'
            )


        document.add_paragraph(
            validation_text
        )


        # =================================================
        # EXPECTED FINAL PATH
        # =================================================

        display_object_path = (
            item.get(
                "display_object_path",
                ""
            )
        )


        if display_object_path:

            p = document.add_paragraph()

            run = p.add_run(
                "Ruta final esperada: "
            )

            run.bold = True


            run = p.add_run(
                display_object_path
            )

            run.font.color.rgb = RGBColor(
                0,
                0,
                255
            )


        # =================================================
        # VALIDATION IMAGE
        # =================================================

        validation_image = (
            render_bip_validation_image(

                object_name=
                    object_name,

                artifact_type=
                    artifact_type,

                folder_name=
                    target_folder_name
            )
        )


        add_centered_image(
            document,
            validation_image,
            Cm(13.5)
        )


        document.add_paragraph("")


# =========================================================
# SPANISH DATE
# =========================================================

def format_spanish_date(
    value
):

    months = {
        1: "Enero",
        2: "Febrero",
        3: "Marzo",
        4: "Abril",
        5: "Mayo",
        6: "Junio",
        7: "Julio",
        8: "Agosto",
        9: "Septiembre",
        10: "Octubre",
        11: "Noviembre",
        12: "Diciembre"
    }

    return (
        f"{months[value.month]} "
        f"{value.day}, "
        f"{value.year}"
    )


# =========================================================
# REMOVE TABLE BORDERS
# =========================================================

def remove_table_borders(
    table
):

    for row in table.rows:

        for cell in row.cells:

            tc_pr = (
                cell._tc
                .get_or_add_tcPr()
            )

            borders = parse_xml(r'''
                <w:tcBorders
                    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">

                    <w:top w:val="nil"/>
                    <w:left w:val="nil"/>
                    <w:bottom w:val="nil"/>
                    <w:right w:val="nil"/>
                    <w:insideH w:val="nil"/>
                    <w:insideV w:val="nil"/>

                </w:tcBorders>
            ''')

            tc_pr.append(
                borders
            )


# =========================================================
# ADD BOTTOM BORDER
# =========================================================

def add_bottom_border(
    cell
):

    tc_pr = (
        cell._tc
        .get_or_add_tcPr()
    )

    borders = parse_xml(r'''
        <w:tcBorders
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">

            <w:bottom
                w:val="single"
                w:sz="8"
                w:space="0"
                w:color="000000"/>

        </w:tcBorders>
    ''')

    tc_pr.append(
        borders
    )


# =========================================================
# COVER PAGE
# =========================================================

def add_im090_cover_page(
    document,
    author_name,
    development_name,
    approvers=None
):

    if approvers is None:

        approvers = [
            "<Aprobador 1>",
            "<Aprobador 2>",
            "<Aprobador 3>"
        ]


    current_dir = os.path.dirname(
        os.path.abspath(
            __file__
        )
    )


    project_root = os.path.dirname(
        os.path.dirname(
            current_dir
        )
    )


    oracle_logo_path = os.path.join(
        project_root,
        "templates",
        "Oracle-Logo.jpg"
    )


    # =====================================================
    # TOP SPACE
    # =====================================================

    for _ in range(3):

        document.add_paragraph("")


    # =====================================================
    # OUM
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        "OUM"
    )

    run.font.name = "Arial"
    run.font.size = Pt(10)


    # =====================================================
    # TITLE
    # =====================================================

    p = document.add_paragraph()

    p.paragraph_format.space_after = Pt(
        2
    )

    run = p.add_run(
        "IM.090 INSTRUCCIONES DE INSTALACIÓN"
    )

    run.font.name = "Arial"
    run.font.size = Pt(17)


    # =====================================================
    # ORACLE TEXT
    # =====================================================

    p = document.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        8
    )

    run = p.add_run(
        "ORACLE"
    )

    run.font.name = "Arial"
    run.font.size = Pt(16)

    run.font.color.rgb = RGBColor(
        0,
        0,
        255
    )


    # =====================================================
    # DEVELOPMENT NAME
    # =====================================================

    title_table = document.add_table(
        rows=1,
        cols=1
    )

    title_table.style = None

    cell = title_table.cell(
        0,
        0
    )

    shading = parse_xml(
        r'<w:shd {} w:fill="F2F2F2"/>'.format(
            nsdecls("w")
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading
    )

    p = cell.paragraphs[0]

    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    run = p.add_run(
        development_name
    )

    run.font.name = "Arial"
    run.font.size = Pt(15)


    document.add_paragraph("")
    document.add_paragraph("")


    # =====================================================
    # INFORMATION
    # =====================================================

    info = document.add_table(
        rows=5,
        cols=2
    )

    info.style = None

    info.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    remove_table_borders(
        info
    )


    now = datetime.today()


    info_data = [

        (
            "Autor:",
            author_name
        ),

        (
            "Fecha de Creación:",
            format_spanish_date(
                now
            )
        ),

        (
            "Última Actualización:",
            now.strftime(
                "%d/%m/%Y"
            )
        ),

        (
            "Document Reference:",
            ""
        ),

        (
            "Versión:",
            "1.0"
        )
    ]


    for index, item in enumerate(
        info_data
    ):

        left = info.cell(
            index,
            0
        )

        right = info.cell(
            index,
            1
        )

        left.text = item[0]
        right.text = item[1]

        left.width = Cm(4)
        right.width = Cm(10)

        for cell_item in [
            left,
            right
        ]:

            for paragraph in (
                cell_item.paragraphs
            ):

                for run in (
                    paragraph.runs
                ):

                    run.font.name = "Arial"
                    run.font.size = Pt(8)


    # =====================================================
    # APPROVERS
    # =====================================================

    document.add_paragraph("")

    p = document.add_paragraph()

    run = p.add_run(
        "Aprobadores:"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)


    approvers_table = document.add_table(
        rows=3,
        cols=2
    )

    approvers_table.style = None

    remove_table_borders(
        approvers_table
    )


    for index in range(3):

        left = approvers_table.cell(
            index,
            0
        )

        right = approvers_table.cell(
            index,
            1
        )

        left.text = (
            approvers[index]
            if index < len(
                approvers
            )
            else ""
        )

        add_bottom_border(
            right
        )

        for paragraph in (
            left.paragraphs
        ):

            paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT
            )

            for run in paragraph.runs:

                run.font.name = "Arial"
                run.font.size = Pt(8)


    # =====================================================
    # ORACLE LOGO
    # =====================================================

    document.add_paragraph("")
    document.add_paragraph("")

    if os.path.exists(
        oracle_logo_path
    ):

        p = document.add_paragraph()

        run = p.add_run()

        run.add_picture(
            oracle_logo_path,
            width=Cm(2.7)
        )


    document.add_page_break()


# =========================================================
# CONTROL DOCUMENT PAGE
# =========================================================

def add_control_document_page(
    document,
    author_name,
    reviewers=None
):

    if reviewers is None:

        reviewers = []


    # =====================================================
    # 1
    # =====================================================

    create_header(
        document,
        "1\tControl del Documento"
    )


    # =====================================================
    # 1.1
    # =====================================================

    create_header(
        document,
        "1.1\tHistorial de Cambios",
        size=13
    )


    history_table = document.add_table(
        rows=2,
        cols=4
    )

    history_table.style = (
        "Table Grid"
    )

    history_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )


    headers = [

        "Fecha",

        "Autor",

        "Versión",

        "Detalle de Cambio"
    ]


    for index, header in enumerate(
        headers
    ):

        cell = history_table.cell(
            0,
            index
        )

        cell.text = header

        apply_table_header_style(
            cell
        )


    today = datetime.today()


    history_table.cell(
        1,
        0
    ).text = today.strftime(
        "%d/%m/%Y"
    )


    history_table.cell(
        1,
        1
    ).text = author_name


    history_table.cell(
        1,
        2
    ).text = "1.0"


    history_table.cell(
        1,
        3
    ).text = "Creación del Documento"


    document.add_paragraph("")
    document.add_paragraph("")


    # =====================================================
    # 1.2
    # =====================================================

    create_header(
        document,
        "1.2\tRevisores",
        size=13
    )


    reviewers_table = document.add_table(
        rows=5,
        cols=2
    )

    reviewers_table.style = (
        "Table Grid"
    )

    reviewers_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )


    reviewers_table.cell(
        0,
        0
    ).text = "Nombre"


    reviewers_table.cell(
        0,
        1
    ).text = "Puesto"


    apply_table_header_style(
        reviewers_table.cell(
            0,
            0
        )
    )


    apply_table_header_style(
        reviewers_table.cell(
            0,
            1
        )
    )


    for index, reviewer in enumerate(
        reviewers[:4],
        start=1
    ):

        reviewers_table.cell(
            index,
            0
        ).text = reviewer.get(
            "name",
            ""
        )

        reviewers_table.cell(
            index,
            1
        ).text = reviewer.get(
            "job",
            ""
        )


    document.add_page_break()


# =========================================================
# INSTALLATION REQUIREMENTS
# =========================================================

def add_installation_requirements(
    document,
    selected_components,
    erp_roles
):

    selected_components = (
        selected_components
        or
        []
    )

    erp_roles = (
        erp_roles
        or
        []
    )


    create_header(
        document,
        "2\tRequerimientos de Instalación"
    )


    # =====================================================
    # ERP CLOUD
    # =====================================================

    if "BI Publisher" in selected_components:

        p = document.add_paragraph(
            style="List Bullet"
        )

        run = p.add_run(
            "Oracle ERP Cloud"
        )

        run.bold = True


        p = document.add_paragraph()

        p.paragraph_format.left_indent = Cm(
            0.75
        )

        p.add_run(
            "Editor de BI de aplicaciones en la nube: "
            "necesario para generar reportes."
        )


        if erp_roles:

            p = document.add_paragraph()

            p.paragraph_format.left_indent = Cm(
                0.75
            )

            p.add_run(
                "El usuario de ERP debe tener asignados "
                "los siguientes roles:"
            )


            for role in erp_roles:

                p = document.add_paragraph()

                p.paragraph_format.left_indent = Cm(
                    1
                )

                p.add_run(
                    role
                )


    # =====================================================
    # PAAS
    # =====================================================

    paas_components = []


    if "OIC" in selected_components:

        paas_components.append(
            (
                "Oracle Integration Cloud (OIC)",
                "necesario para la creación de integraciones."
            )
        )


    if "Visual Builder" in selected_components:

        paas_components.append(
            (
                "Visual Builder (VBCS)",
                "necesario para la construcción de la interfaz visual."
            )
        )


    if "Objetos BD" in selected_components:

        paas_components.append(
            (
                "Oracle Database",
                "base de datos Oracle."
            )
        )


    if paas_components:

        document.add_paragraph("")

        p = document.add_paragraph(
            style="List Bullet"
        )

        run = p.add_run(
            "Platform as a Service - PaaS "
            "(Plataforma como Servicio)"
        )

        run.bold = True


        for component_name, description in (
            paas_components
        ):

            p = document.add_paragraph()

            p.paragraph_format.left_indent = Cm(
                0.75
            )

            run = p.add_run(
                f"{component_name}: "
            )

            run.bold = False

            p.add_run(
                description
            )


# =========================================================
# ADD INFORMATION BOX
# =========================================================

def add_information_box(
    document,
    title,
    text,
    fill="FFF2CC"
):

    table = document.add_table(
        rows=1,
        cols=1
    )

    table.autofit = False

    cell = table.cell(
        0,
        0
    )

    shading = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(
            nsdecls("w"),
            fill
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading
    )


    paragraph = cell.paragraphs[0]

    paragraph.paragraph_format.space_after = Pt(
        0
    )


    run = paragraph.add_run(
        f"{title}: "
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


    run = paragraph.add_run(
        text
    )

    run.font.name = "Arial"
    run.font.size = Pt(9)


# =========================================================
# ADD SQL CODE BLOCK
# =========================================================

def add_sql_code_block(
    document,
    sql_text
):

    table = document.add_table(
        rows=1,
        cols=1
    )

    table.autofit = False

    cell = table.cell(
        0,
        0
    )


    shading = parse_xml(
        r'<w:shd {} w:fill="F2F2F2"/>'.format(
            nsdecls("w")
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading
    )


    paragraph = cell.paragraphs[
        0
    ]

    paragraph.paragraph_format.space_before = Pt(
        3
    )

    paragraph.paragraph_format.space_after = Pt(
        3
    )


    run = paragraph.add_run(
        sql_text
    )

    run.font.name = "Courier New"
    run.font.size = Pt(8)


# =========================================================
# FORMAT INSTALLATION OBJECTS
# =========================================================

def format_installation_objects(
    objects
):

    if not objects:

        return "-"


    clean_objects = [

        str(value).strip()

        for value in objects

        if str(value).strip()
    ]


    if not clean_objects:

        return "-"


    return ", ".join(
        clean_objects
    )


# =========================================================
# ADD INSTALLATION SCRIPTS TABLE
# =========================================================

def add_installation_scripts_table(
    document,
    installation_scripts
):

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = (
        "Table Grid"
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )


    headers = [
        "Orden",
        "Tipo",
        "Archivo",
        "Objetos Incluidos"
    ]


    for index, header in enumerate(
        headers
    ):

        cell = table.cell(
            0,
            index
        )

        cell.text = header

        apply_table_header_style(
            cell
        )


    for script in installation_scripts:

        cells = table.add_row().cells


        cells[0].text = str(
            script.get(
                "order",
                ""
            )
        )


        cells[1].text = str(
            script.get(
                "type",
                ""
            )
        )


        cells[2].text = str(
            script.get(
                "file_name",
                ""
            )
        )


        cells[3].text = (
            format_installation_objects(
                script.get(
                    "objects",
                    []
                )
            )
        )


        for cell in cells:

            for paragraph in (
                cell.paragraphs
            ):

                for run in (
                    paragraph.runs
                ):

                    run.font.name = "Arial"
                    run.font.size = Pt(8)


    # =====================================================
    # WIDTHS
    # =====================================================

    for row in table.rows:

        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(3.0)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(7.0)


    return table

# =========================================================
# DATABASE INSTALLATION SECTION
# =========================================================

def add_database_installation_section(
    document,
    schema_name,
    database_export_info
):

    database_export_info = (
        database_export_info
        or
        {}
    )


    installation_scripts = (
        database_export_info.get(
            "installation_scripts",
            []
        )
    )


    schema_name = (
        schema_name
        or
        "<ESQUEMA>"
    )


    # =====================================================
    # NEW PAGE
    # =====================================================

    document.add_page_break()


    # =====================================================
    # 3
    # =====================================================

    create_header(
        document,
        "3\tInstrucciones para la instalación del Módulo"
    )


    # =====================================================
    # 3.1
    # =====================================================

    create_header(
        document,
        "3.1\tInstalación paso a paso de Objetos de Base de Datos",
        size=13
    )


    # =====================================================
    # INTRODUCTION
    # =====================================================

    paragraph = document.add_paragraph()

    paragraph.add_run(
        "En esta sección se detallan las actividades "
        "necesarias para realizar la instalación de los "
        "objetos de Base de Datos incluidos en el paquete "
        "de despliegue. Los scripts deberán ejecutarse "
        "respetando estrictamente el orden indicado en "
        "el presente documento."
    )


    # =====================================================
    # 3.1.1
    # =====================================================

    create_header(
        document,
        "3.1.1\tPreparación y conexión a la Base de Datos",
        size=11
    )


    preparation_steps = [

        (
            "Abrir Oracle SQL Developer o la herramienta "
            "autorizada para la administración de la "
            "Base de Datos Oracle."
        ),

        (
            "Ubicar la conexión correspondiente al "
            "ambiente destino donde se realizará la "
            "instalación."
        ),

        (
            "Hacer clic derecho sobre la conexión y "
            "seleccionar la opción Conectar."
        ),

        (
            "Ingresar la contraseña correspondiente "
            "al usuario configurado."
        ),

        (
            "No modificar el usuario definido en la "
            "conexión, salvo que exista una instrucción "
            "expresa dentro del procedimiento de pase."
        ),

        (
            "Confirmar que la conexión se encuentre "
            "activa antes de continuar con la ejecución "
            "de los scripts."
        )
    ]


    for index, step in enumerate(
        preparation_steps,
        start=1
    ):

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.left_indent = Cm(
            0.5
        )

        run = paragraph.add_run(
            f"{index}. "
        )

        run.bold = True

        paragraph.add_run(
            step
        )


    add_information_box(

        document,

        "Importante",

        (
            "Antes de iniciar la instalación, valide que "
            "la conexión corresponda al ambiente destino "
            "definido para el pase."
        ),

        fill="D9EAF7"
    )


    # =====================================================
    # 3.1.2
    # =====================================================

    create_header(
        document,
        "3.1.2\tValidación del esquema de instalación",
        size=11
    )


    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        "Esquema de instalación: "
    )

    run.bold = True


    run = paragraph.add_run(
        schema_name
    )

    run.bold = True


    add_information_box(

        document,

        "Importante",

        (
            "Todos los scripts deberán ejecutarse bajo "
            f"el esquema {schema_name}. La ejecución "
            "sobre un esquema diferente podría generar "
            "objetos o información en una ubicación "
            "incorrecta."
        )
    )


    document.add_paragraph("")


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "Antes de ejecutar los scripts, valide el "
        "esquema actual de la sesión mediante la "
        "siguiente sentencia:"
    )


    document.add_paragraph("")


    add_sql_code_block(
        document,
        (
            "SELECT SYS_CONTEXT(\n"
            "           'USERENV',\n"
            "           'CURRENT_SCHEMA'\n"
            "       ) AS CURRENT_SCHEMA\n"
            "FROM DUAL;"
        )
    )


    document.add_paragraph("")


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "El resultado esperado debe corresponder al "
        "siguiente esquema:"
    )


    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = Cm(
        0.75
    )

    run = paragraph.add_run(
        schema_name
    )

    run.bold = True


    document.add_paragraph("")


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "Cuando la sesión se encuentre autenticada con "
        "un usuario diferente y dicho usuario disponga "
        "de los privilegios necesarios sobre el esquema "
        "destino, se podrá establecer el esquema de "
        "trabajo mediante:"
    )


    document.add_paragraph("")


    add_sql_code_block(
        document,
        (
            "ALTER SESSION SET CURRENT_SCHEMA = "
            f"{schema_name};"
        )
    )


    document.add_paragraph("")


    add_information_box(

        document,

        "Nota",

        (
            "ALTER SESSION SET CURRENT_SCHEMA modifica "
            "el esquema utilizado para la resolución de "
            "objetos durante la sesión. Esta instrucción "
            "no modifica el usuario autenticado ni otorga "
            "privilegios adicionales."
        ),

        fill="E2F0D9"
    )


    # =====================================================
    # 3.1.3
    # =====================================================

    create_header(
        document,
        "3.1.3\tEjecución de scripts",
        size=11
    )


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "Los archivos incluidos en la carpeta "
    )


    run = paragraph.add_run(
        "scripts"
    )

    run.bold = True


    paragraph.add_run(
        " deberán ejecutarse en el orden establecido "
        "en la siguiente tabla:"
    )


    document.add_paragraph("")


    if installation_scripts:

        add_installation_scripts_table(

            document,

            installation_scripts
        )

    else:

        add_information_box(

            document,

            "Advertencia",

            (
                "No se encontraron scripts de Base de "
                "Datos para incluir en el plan de "
                "instalación."
            ),

            fill="FCE4D6"
        )


    document.add_paragraph("")


    # =====================================================
    # GENERAL EXECUTION INSTRUCTIONS
    # =====================================================

    execution_steps = [

        (
            "Ubicar la carpeta scripts incluida en "
            "el paquete de instalación."
        ),

        (
            "Abrir el primer archivo indicado en la "
            "tabla de orden de ejecución."
        ),

        (
            "Cargar el contenido del archivo en una "
            "hoja de trabajo de Oracle SQL Developer."
        ),

        (
            "Ejecutar el archivo utilizando la opción "
            "Run Script o la tecla F5."
        ),

        (
            "Revisar la salida de ejecución y confirmar "
            "que el script haya finalizado sin errores."
        ),

        (
            "Continuar con el siguiente archivo únicamente "
            "cuando el script anterior haya finalizado "
            "correctamente."
        ),

        (
            "Repetir el procedimiento hasta completar "
            "todos los archivos definidos en el plan de "
            "instalación."
        )
    ]


    for index, step in enumerate(
        execution_steps,
        start=1
    ):

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.left_indent = Cm(
            0.5
        )

        run = paragraph.add_run(
            f"{index}. "
        )

        run.bold = True

        paragraph.add_run(
            step
        )


    # =====================================================
    # INDIVIDUAL SCRIPT ORDER
    # =====================================================

    if installation_scripts:

        document.add_paragraph("")


        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            "Orden de ejecución:"
        )

        run.bold = True


        for script in installation_scripts:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.left_indent = Cm(
                0.75
            )


            order = script.get(
                "order",
                ""
            )


            file_name = script.get(
                "file_name",
                ""
            )


            script_type = script.get(
                "type",
                ""
            )


            run = paragraph.add_run(
                f"{order}. "
            )

            run.bold = True


            run = paragraph.add_run(
                file_name
            )

            run.bold = True


            paragraph.add_run(
                f" — {script_type}"
            )


    add_information_box(

        document,

        "Importante",

        (
            "No alterar el orden de ejecución establecido. "
            "Los archivos han sido organizados considerando "
            "las dependencias técnicas entre los diferentes "
            "tipos de objetos de Base de Datos."
        )
    )


    # =====================================================
    # 3.1.4
    # =====================================================

    create_header(
        document,
        "3.1.4\tValidación posterior a la instalación",
        size=11
    )


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "Una vez ejecutados todos los scripts, se deberá "
        "validar que los objetos hayan sido creados "
        "correctamente y que no existan objetos inválidos "
        "asociados a la instalación."
    )


    document.add_paragraph("")


    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        "Validación de objetos inválidos"
    )

    run.bold = True


    document.add_paragraph("")


    add_sql_code_block(
        document,
        (
            "SELECT\n"
            "    OBJECT_NAME,\n"
            "    OBJECT_TYPE,\n"
            "    STATUS\n"
            "FROM USER_OBJECTS\n"
            "WHERE STATUS = 'INVALID'\n"
            "ORDER BY\n"
            "    OBJECT_TYPE,\n"
            "    OBJECT_NAME;"
        )
    )


    document.add_paragraph("")


    paragraph = document.add_paragraph()

    paragraph.add_run(
        "Si existen objetos inválidos, revisar los "
        "errores de compilación mediante:"
    )


    document.add_paragraph("")


    add_sql_code_block(
        document,
        (
            "SELECT\n"
            "    NAME,\n"
            "    TYPE,\n"
            "    LINE,\n"
            "    POSITION,\n"
            "    TEXT\n"
            "FROM USER_ERRORS\n"
            "ORDER BY\n"
            "    NAME,\n"
            "    SEQUENCE;"
        )
    )


    document.add_paragraph("")


    add_information_box(

        document,

        "Resultado esperado",

        (
            "Todos los scripts deben finalizar sin errores "
            "y los objetos instalados deberán encontrarse "
            "en estado VALID."
        ),

        fill="E2F0D9"
    )


# =========================================================
# GENERATE INSTALLATION MANUAL
# =========================================================

def generate_installation_manual(
    author_name,
    development_name,
    selected_components,
    erp_roles=None,
    reviewers=None,
    approvers=None,
    schema_name="",
    database_export_info=None,
    bip_installation_plan=None
):

    document = Document()


    # =====================================================
    # UPDATE WORD FIELDS
    # =====================================================

    settings = (
        document.settings.element
    )


    existing_update_fields = (
        settings.find(
            qn(
                "w:updateFields"
            )
        )
    )


    if existing_update_fields is not None:

        settings.remove(
            existing_update_fields
        )


    update_fields = OxmlElement(
        "w:updateFields"
    )

    update_fields.set(
        qn("w:val"),
        "1"
    )

    settings.append(
        update_fields
    )


    # =====================================================
    # STYLES
    # =====================================================

    create_document_styles(
        document
    )


    normal = document.styles[
        "Normal"
    ]

    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Arial"
    )


    # =====================================================
    # PAGE 1
    # =====================================================

    add_im090_cover_page(

        document,

        author_name,

        development_name,

        approvers
    )


    # =====================================================
    # PAGE 2
    # =====================================================

    add_control_document_page(

        document,

        author_name,

        reviewers
    )


    # =====================================================
    # PAGE 3 - TABLE OF CONTENTS
    # =====================================================

    toc_placeholder = (
        create_toc_table(
            document
        )
    )


    document.add_page_break()


    # =====================================================
    # PAGE 4 - INSTALLATION REQUIREMENTS
    # =====================================================

    add_installation_requirements(

        document,

        selected_components,

        erp_roles
    )

    # =====================================================
    # PAGE 5+ - DATABASE INSTALLATION
    # =====================================================

    if "Objetos BD" in selected_components:

        add_database_installation_section(

            document,

            schema_name,

            database_export_info
        )

    if bip_installation_plan:

        document.add_page_break()

        add_bip_installation_section(
            document,
            bip_installation_plan
        )


    # =====================================================
    # POPULATE TOC AFTER ALL HEADERS EXIST
    # =====================================================

    populate_toc_table(
        document,
        toc_placeholder
    )


    # =====================================================
    # OUTPUT
    # =====================================================

    output = BytesIO()

    document.save(
        output
    )

    output.seek(
        0
    )


    return output