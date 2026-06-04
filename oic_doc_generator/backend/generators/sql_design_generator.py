# =========================================================
# FILE:
# oic_doc_generator/generators/sql_design_generator.py
# =========================================================

from docx.shared import (
    Pt,
    RGBColor
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from oic_doc_generator.backend.utils.word_utils import (
    create_header,
    add_description_box
)

from oic_doc_generator.backend.utils.sql_formatter import (
    beautify_sql
)

from oic_doc_generator.backend.utils.xsd_formatter import (
    build_xsd_text
)

from oic_doc_generator.backend.utils.code_block_utils import (
    add_code_block
)

# =========================================================
# ADD DATASET SECTION
# =========================================================

def add_dataset_section(
    document,
    dataset
):

    dataset_name = dataset.get(
        "name",
        "Dataset"
    )

    sql_text = dataset.get(
        "sql",
        ""
    )

    # =====================================================
    # TITLE
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        f"• Data Set: {dataset_name}"
    )

    run.bold = True

    run.font.name = "Arial"

    run.font.size = Pt(10)

    # =====================================================
    # SQL
    # =====================================================

    beautified = beautify_sql(
        sql_text
    )

    add_code_block(
        document,
        beautified,
        language="sql"
    )

    document.add_paragraph("")


# =========================================================
# ADD XSD SECTION
# =========================================================

def add_xsd_section(
    document,
    xsd_structure
):

    p = document.add_paragraph()

    run = p.add_run(
        "• Estructura XSD:"
    )

    run.bold = True

    run.font.name = "Arial"

    run.font.size = Pt(10)

    xsd_text = build_xsd_text(
        xsd_structure
    )

    add_code_block(
        document,
        xsd_text,
        language="xml"
    )

    document.add_paragraph("")


# =========================================================
# ADD SQL DESIGN SECTION
# =========================================================

def add_sql_design_section(
    document,
    bip_metadata
):

    # =====================================================
    # TITLE
    # =====================================================

    create_header(
        document,
        "6\tSentencias SQL"
    )

    # =====================================================
    # DESCRIPTION
    # =====================================================

    add_description_box(

        document,

        "En esta sección se documentan "
        "las consultas SQL de cada uno "
        "de los data set que contienen "
        "los modelos de datos indicados "
        "en la sección de Diseño del Reporte."
    )

    document.add_paragraph("")

    reports = bip_metadata.get(
        "reports",
        []
    )

    if not reports:

        document.add_paragraph(
            "No se encontraron reportes."
        )

        return

    # =====================================================
    # ITERATE REPORTS
    # =====================================================

    counter = 1

    for report in reports:

        report_name = report.get(
            "report_name",
            ""
        )

        xsd_structure = report.get(
            "xsd_structure",
            {}
        )

        datasets = report.get(
            "dataset_sqls",
            []
        )

        # =================================================
        # HEADER
        # =================================================

        create_header(
            document,
            f"6.{counter}\t{report_name}",
            size=13
        )

        # =================================================
        # XSD
        # =================================================

        add_xsd_section(
            document,
            xsd_structure
        )

        # =================================================
        # DATASETS
        # =================================================

        if not datasets:

            document.add_paragraph(
                "No se encontraron Data Sets."
            )

        else:

            for dataset in datasets:

                add_dataset_section(
                    document,
                    dataset
                )

        document.add_paragraph("")
        document.add_paragraph("")

        counter += 1