# =========================================================
# FILE:
# oic_doc_generator/generators/word_generator.py
# =========================================================

from io import BytesIO
from datetime import datetime
import os

from oic_doc_generator.api.job_manager import (
    advance_step,
    update_activity
)

from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from oic_doc_generator.backend.generators.database_design_generator import (
    add_database_design_section
)

from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)

from docx.oxml import parse_xml

from docx.oxml.ns import (
    nsdecls,
    qn
)

from oic_doc_generator.backend.parsers.iar_parser import (
    find_all_iar_files,
    extract_iar
)

from oic_doc_generator.backend.parsers.project_parser import (
    get_project_name,
    read_project_xml,
    build_application_map,
    get_project_root,
    get_project_version
)

from oic_doc_generator.backend.parsers.orchestration_parser import (
    get_endpoint_flows,
    generate_description
)

from oic_doc_generator.backend.parsers.connections_parser import (
    analyze_dbaas_jca,
    find_application_folder,
    build_action_description
)

from oic_doc_generator.backend.parsers.schedule_parser import (
    get_schedule_information,
    is_scheduled_integration
)

from oic_doc_generator.backend.parsers.service_parser import (
    get_request_response_schemas
)

from oic_doc_generator.backend.generators.service_design_generator import (
    add_service_design_section
)

from oic_doc_generator.backend.parsers.bip_archive_parser import (
    build_bip_artifact_tree
)

from oic_doc_generator.backend.parsers.bip_metadata_builder import (
    build_bip_metadata
)

from oic_doc_generator.backend.generators.report_design_generator import (
    add_report_design_section
)

from oic_doc_generator.backend.generators.sql_design_generator import ( add_sql_design_section )

from oic_doc_generator.backend.utils.xml_utils import (
    clean_tag,
    extract_application_from_refuri
)

from oic_doc_generator.backend.generators.diagram_generator import (
    generate_sequence_diagram_png
)

from oic_doc_generator.backend.parsers.vb_extractor import (
    build_extraction_metadata
)

from oic_doc_generator.backend.parsers.vb_parser import (
    build_page_metadata,
    build_virtual_apps
)

from oic_doc_generator.backend.renderers.html_builder import (
    build_complete_html
)

from oic_doc_generator.backend.renderers.screenshot_renderer import (
    render_html_to_image
)

from oic_doc_generator.backend.utils.word_utils import (
    create_header,
    add_description_box
)

# =========================================================
# APPLY HEADER STYLE
# =========================================================

def apply_header_style(
    cell,
    fill="D9D9D9",
    white_text=False
):

    shading_elm = parse_xml(

        r'<w:shd {} w:fill="{}"/>'.format(
            nsdecls('w'),
            fill
        )
    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.bold = True

            if white_text:

                run.font.color.rgb = RGBColor(
                    255,
                    255,
                    255
                )

# =========================================================
# ADD FRAMED IMAGE
# =========================================================

def add_framed_image(
    document,
    image_path,
    width=Inches(6.5)
):

    table = document.add_table(
        rows=1,
        cols=1
    )

    # =====================================================
    # REMOVE DEFAULT TABLE STYLE
    # =====================================================

    table.style = None

    cell = table.cell(0,0)

    # =====================================================
    # BORDER ONLY FOR IMAGE CELL
    # =====================================================

    tc_pr = cell._tc.get_or_add_tcPr()

    borders = parse_xml(r'''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="20" w:color="000000"/>
            <w:left w:val="single" w:sz="20" w:color="000000"/>
            <w:bottom w:val="single" w:sz="20" w:color="000000"/>
            <w:right w:val="single" w:sz="20" w:color="000000"/>
        </w:tcBorders>
    ''')

    tc_pr.append(
        borders
    )

    paragraph = cell.paragraphs[0]

    paragraph.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    run = paragraph.add_run()

    run.add_picture(
        image_path,
        width=width
    )

# =========================================================
# CREATE CUSTOM STYLES
# =========================================================

def create_custom_styles(
    document
):

    styles = document.styles

    # =====================================================
    # HD1
    # =====================================================

    if "HD1" not in styles:

        style = styles.add_style(
            "HD1",
            WD_STYLE_TYPE.PARAGRAPH
        )

        font = style.font

        font.name = "Arial"
        font.size = Pt(16)
        font.bold = True

        paragraph_format = style.paragraph_format

        paragraph_format.left_indent = Cm(0)

        paragraph_format.first_line_indent = Cm(0)

        paragraph_format.space_before = Pt(6)

        paragraph_format.space_after = Pt(12)

        paragraph_format.keep_with_next = True

        paragraph_format.keep_together = True

        # ================================================
        # BORDER TOP
        # ================================================

        pPr = style.element.get_or_add_pPr()

        border_xml = parse_xml(r'''
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top
                    w:val="single"
                    w:sz="45"
                    w:space="1"
                    w:color="auto"/>
            </w:pBdr>
        ''')

        pPr.append(border_xml)

    # =====================================================
    # HD2
    # =====================================================

    if "HD2" not in styles:

        style = styles.add_style(
            "HD2",
            WD_STYLE_TYPE.PARAGRAPH
        )

        font = style.font

        font.name = "Arial"
        font.size = Pt(13)
        font.bold = True

        # ================================================
        # BORDER TOP
        # ================================================

        pPr = style.element.get_or_add_pPr()

        border_xml = parse_xml(r'''
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top
                    w:val="single"
                    w:sz="24"
                    w:space="1"
                    w:color="auto"/>
            </w:pBdr>
        ''')

        pPr.append(border_xml)

        paragraph_format = style.paragraph_format

        paragraph_format.left_indent = Cm(0)

        paragraph_format.first_line_indent = Cm(0)

        paragraph_format.space_before = Pt(6)

        paragraph_format.space_after = Pt(6)

        paragraph_format.keep_with_next = True

        paragraph_format.keep_together = True

# =========================================================
# ADD COVER PAGE
# =========================================================

def add_cover_page(
    document,
    development_name,
    author_name
):

    current_dir = os.path.dirname(
        os.path.abspath(__file__)
    )

    project_root = os.path.dirname(
        os.path.dirname(current_dir)
    )

    logo_path = os.path.join(

        project_root,

        "templates",

        "NEORA_PNG.png"
    )

    # =====================================================
    # HEADER TABLE
    # =====================================================

    table = document.add_table(
        rows=4,
        cols=4
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.style = "Table Grid"

    # =====================================================
    # COLUMN WIDTHS
    # =====================================================

    widths = [

        Cm(5),

        Cm(9),

        Cm(4),

        Cm(7)
    ]

    for row in table.rows:

        for idx, width in enumerate(widths):

            row.cells[idx].width = width

    # =====================================================
    # MERGES
    # =====================================================

    logo_cell = table.cell(0,0).merge(
        table.cell(3,0)
    )

    title_cell = table.cell(0,1).merge(
        table.cell(3,1)
    )

    # =====================================================
    # LOGO
    # =====================================================

    p = logo_cell.paragraphs[0]

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run()

    if os.path.exists(logo_path):

        run.add_picture(
            logo_path,
            width=Cm(4)
        )

    # =====================================================
    # TITLE
    # =====================================================

    p = title_cell.paragraphs[0]

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(
        "ESPECIFICACIÓN\nDE\nDESARROLLO"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    # =====================================================
    # GREEN COLUMN
    # =====================================================

    labels = [

        "Código",

        "Versión",

        "Fecha",

        "Página"
    ]

    values = [

        "NEO-GD-IN-02",

        "00",

        "15/06/2023",

        "1 de 1"
    ]

    for i in range(4):

        left = table.cell(i,2)
        right = table.cell(i,3)

        left.text = labels[i]
        right.text = values[i]

        apply_header_style(
            left,
            fill="6DBE45",
            white_text=True
        )

        for p in left.paragraphs:

            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

        for p in right.paragraphs:

            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

    # =====================================================
    # SPACE
    # =====================================================

    document.add_paragraph("")

    # =====================================================
    # BLACK LINE
    # =====================================================

    p = document.add_paragraph()

    run = p.add_run(
        " "
    )

    run.font.highlight_color = None

    p.paragraph_format.space_after = Pt(0)

    shading = parse_xml(

        r'<w:shd {} w:fill="000000"/>'.format(
            nsdecls('w')
        )
    )

    p._p.get_or_add_pPr().append(
        shading
    )

    # =====================================================
    # CONTENT
    # =====================================================

    def add_line(
        text,
        size,
        bold=False,
        color=None
    ):

        p = document.add_paragraph()

        p.alignment = (
            WD_PARAGRAPH_ALIGNMENT.LEFT
        )

        run = p.add_run(text)

        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold

        if color:

            run.font.color.rgb = color

    add_line(
        "OUM",
        18
    )

    add_line(
        "DS.140.1 ESPECIFICACIÓN DE DISEÑO",
        24
    )

    add_line(
        "NEORA",
        24,
        color=RGBColor(0,0,255)
    )

    add_line(
        development_name,
        24
    )

    # =====================================================
    # INFO TABLE
    # =====================================================

    info = document.add_table(
    rows=5,
    cols=2
)

    info.style = "Table Grid"

    # =====================================================
    # REMOVE BORDERS
    # =====================================================

    for row in info.rows:

        for cell in row.cells:

            tc_pr = cell._tc.get_or_add_tcPr()

            borders = parse_xml(r'''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="nil"/>
                    <w:left w:val="nil"/>
                    <w:bottom w:val="nil"/>
                    <w:right w:val="nil"/>
                </w:tcBorders>
            ''')

            tc_pr.append(borders)

    info_data = [

        ("Autor:", author_name),

        (
            "Fecha de Creación:",
            datetime.today().strftime(
                "%B %d, %Y"
            )
        ),

        (
            "Última Actualización:",
            datetime.today().strftime(
                "%B %d, %Y"
            )
        ),

        ("Referencia:", ""),

        ("Versión:", "1.0")
    ]

    info_data = [

        ("Autor:", author_name),

        (
            "Fecha de Creación:",
            datetime.today().strftime(
                "%B %d, %Y"
            )
        ),

        (
            "Última Actualización:",
            datetime.today().strftime(
                "%B %d, %Y"
            )
        ),

        ("Referencia:", ""),

        ("Versión:", "1.0")
    ]

    for idx, item in enumerate(info_data):

        left = info.cell(idx,0)
        right = info.cell(idx,1)

        left.text = item[0]
        right.text = item[1]

        for cell in [left, right]:

            for p in cell.paragraphs:

                for run in p.runs:

                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    # =====================================================
    # SPACES
    # =====================================================

    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph()

    run = p.add_run(
        "Aprobadores:"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    # =====================================================
    # BOTTOM LOGO
    # =====================================================

    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")

    if os.path.exists(logo_path):

        p = document.add_paragraph()

        run = p.add_run()

        run.add_picture(
            logo_path,
            width=Cm(4)
        )

    # =====================================================
    # PAGE BREAK
    # =====================================================

    document.add_page_break()



# =========================================================
# FORMAT JSON WRAPPER
# =========================================================

def format_json_wrapper(
    elements,
    is_array=False
):

    if not elements:

        return "{}"

    lines = []

    if is_array:

        lines.append("[")
        lines.append("{")

    else:

        lines.append("{")

    for element in elements:

        name = element.get(
            "name",
            ""
        )

        data_type = element.get(
            "type",
            "string"
        )

        if name in [

            "response-wrapper",

            "request-wrapper",

            "topLevelArray"
        ]:

            continue

        lines.append(
            f'  "{name}": "{data_type}",'
        )

    if len(lines) > 1:

        lines[-1] = lines[-1].rstrip(",")

    if is_array:

        lines.append("}")
        lines.append("]")

    else:

        lines.append("}")

    return "\n".join(lines)


# =========================================================
# ADD PROCESS TABLE
# =========================================================

def add_process_table(
    document,
    rows
):

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.style = "Table Grid"

    headers = [

        "Nro. Secuencia",

        "Nombre Acción",

        "Descripción de la Acción"
    ]

    hdr = table.rows[0].cells

    for i in range(3):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )

    for row in rows:

        cells = table.add_row().cells

        cells[0].text = str(
            row["Nro. Secuencia"]
        )

        cells[1].text = (
            row["Nombre Acción"]
        )

        cells[2].text = (
            row["Descripción de la Acción"]
        )


# =========================================================
# GENERATE WORD DOCUMENT
# =========================================================

def generate_word_document(
    package_path,
    author_name,
    development_name,
    selected_components,
    visual_builder_apps,
    apex_apps,
    bip_files=None,
    database_metadata=None,
    database_export_info=None,
    job_id=None
):

    document = Document()

    create_custom_styles(
        document
    )

    add_cover_page(
        document,
        development_name,
        author_name
    )

    style = document.styles['Normal']

    style.font.name = 'Arial'

    style._element.rPr.rFonts.set(

        qn('w:eastAsia'),

        'Arial'
    )

    style.font.size = Pt(10)

    # =====================================================
    # 1 CONTROL DOCUMENTO
    # =====================================================

    create_header(
        document,
        "1\tControl del Documento"
    )

    create_header(
        document,
        "1.1\tRegistro de Modificaciones"
    )

    table = document.add_table(
        rows=2,
        cols=4
    )

    table.style = "Table Grid"

    headers = [

        "Fecha",

        "Autor",

        "Versión",

        "Descripción del Cambio"
    ]

    hdr = table.rows[0].cells

    for i in range(4):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )

    today = datetime.today().strftime(
        "%d/%m/%Y"
    )

    row = table.rows[1].cells

    row[0].text = today
    row[1].text = author_name
    row[2].text = "1.0"
    row[3].text = (
        "Creación del Documento"
    )

    document.add_paragraph("")

    # =====================================================
    # 1.2 REVISORES
    # =====================================================

    create_header(
        document,
        "1.2\tRevisores"
    )

    reviewers_table = document.add_table(
        rows=5,
        cols=2
    )

    reviewers_table.style = "Table Grid"

    hdr = reviewers_table.rows[0].cells

    hdr[0].text = "Nombre"
    hdr[1].text = "Puesto"

    apply_header_style(
        hdr[0]
    )

    apply_header_style(
        hdr[1]
    )

    document.add_page_break()

    create_header(
        document,
        "Tabla de Contenido"
    )

    paragraph = document.add_paragraph()

    run = paragraph.add_run()

    fld_char_begin = parse_xml(
        r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    instr_text = parse_xml(
        r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> TOC \o "1-3" \h \z \u </w:instrText>'
    )

    fld_char_separate = parse_xml(
        r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    fld_char_end = parse_xml(
        r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    run._r.append(
        fld_char_begin
    )

    run._r.append(
        instr_text
    )

    run._r.append(
        fld_char_separate
    )

    run._r.append(
        fld_char_end
    )

    document.add_page_break()

    # =====================================================
    # 2 VISIÓN GENERAL
    # =====================================================
    document.add_page_break()
    create_header(
        document,
        "2\tVisión General"
    )

    add_description_box(

        document,

        "Descripción general del desarrollo "
        "que permite identificar los componentes "
        "y la manera que estos interactúan "
        "como parte del desarrollo."
    )

    document.add_paragraph("")

    document.add_paragraph(

        f"En este documento se tiene "
        f"como propósito describir "
        f"cada componente del "
        f"{development_name} y su "
        f"funcionalidad a nivel "
        f"técnico-funcional."
    )

    # =====================================================
    # 2.1
    # =====================================================

    create_header(
        document,
        "2.1\tReglas del Negocio"
    )

    add_description_box(

        document,

        "En esta sección se coloca la información "
        "de reglas de negocio descritas en el "
        "documento de análisis funcional del "
        "desarrollo (AN-100) como parte de la "
        "documentación se revisa que todo lo que "
        "se coloca en esta sección se cumple "
        "por el desarrollo."
    )

    document.add_paragraph("")

    for _ in range(4):

        document.add_paragraph("")

    # =====================================================
    # 2.2
    # =====================================================

    create_header(
        document,
        "2.2\tListado de Componentes"
    )

    document.add_paragraph(

        "El detalle de los "
        "componentes necesarios "
        "para el presente desarrollo "
        "se lista a continuación:"
    )

    components_table = document.add_table(
        rows=1,
        cols=2
    )

    components_table.style = "Table Grid"

    hdr = components_table.rows[0].cells

    hdr[0].text = "Componente"
    hdr[1].text = "Descripción"

    apply_header_style(
        hdr[0],
        fill="808080",
        white_text=True
    )

    apply_header_style(
        hdr[1],
        fill="808080",
        white_text=True
    )

    for component in selected_components:

        row = (
            components_table.add_row().cells
        )

        row[0].text = component
        row[1].text = ""

    document.add_paragraph("")

    # =====================================================
    # 2.3
    # =====================================================

    create_header(
        document,
        "2.3\tDiagrama de componentes"
    )

    document.add_paragraph(

        "El siguiente diagrama "
        "representa los componentes "
        "y la relación existente "
        "entre ellos."
    )

    document.add_paragraph("")
    document.add_paragraph("")

    # =====================================================
    # 2.4
    # =====================================================

    create_header(
        document,
        "2.4\tDiagrama de Secuencia"
    )

    document.add_paragraph(

        "El siguiente diagrama "
        "representa la secuencia "
        "que existe entre los "
        "componentes involucrados:"
    )

    document.add_paragraph("")
    document.add_paragraph("")


    # =====================================================
    # 2.5 OIC
    # =====================================================

    create_header(
        document,
        "2.5\tDiagrama de bloques de construcción"
    )

    add_description_box(

        document,

        "Se presentan imágenes de cada integración "
        "creada para el desarrollo y su respectivo "
        "diagrama de secuencia."
    )

    document.add_paragraph("")

    integrations_data = []

    if not package_path:

        document.add_paragraph(
            "No aplica."
        )

    else:

        iar_files = find_all_iar_files(
            package_path
        )

        for iar in iar_files:

            extracted_iar = extract_iar(
                iar
            )

            integration_name = (
                get_project_name(
                    extracted_iar
                )
            )

            applications = (
                read_project_xml(
                    extracted_iar
                )
            )

            app_map = build_application_map(
                applications
            )

            root = get_project_root(
                extracted_iar
            )

            integrations_data.append({

                "name":
                    integration_name,

                "path":
                    extracted_iar,

                "scheduled":
                    is_scheduled_integration(
                        extracted_iar
                    ),

                "version":
                    get_project_version(
                        extracted_iar
                    )
            })

            p = document.add_paragraph()

            run = p.add_run(
                f"• {integration_name}"
            )

            run.bold = True

            endpoint_flows = (
                get_endpoint_flows(
                    root,
                    app_map
                )
            )

            # =============================================
            # ENDPOINTS
            # =============================================

            for endpoint, flow in (

                endpoint_flows.items()
            ):

                description = (
                    generate_description(

                        endpoint,

                        flow,

                        applications,

                        extracted_iar,

                        root
                    )
                )

                document.add_paragraph(
                    description
                )

                document.add_paragraph(

                    "A continuación, "
                    "se describe cada "
                    "una de las acciones "
                    "del flujo de integración:"
                )

                rows = []

                seq = 1

                if is_scheduled_integration(
                    extracted_iar
                ):

                    start_desc = (
                        "Integración programada."
                    )

                else:

                    start_desc = (
                        "De tipo REST, endpoint."
                    )

                rows.append({

                    "Nro. Secuencia":
                        seq,

                    "Nombre Acción":
                        endpoint,

                    "Descripción de la Acción":
                        start_desc
                })

                seq += 1

                # =========================================
                # FLOW ACTIONS
                # =========================================

                for elem in flow.iter():

                    try:

                        tag = clean_tag(
                            elem.tag
                        )

                    except:

                        continue

                    # =====================================
                    # INVOKE
                    # =====================================

                    if tag == "invoke":

                        refuri = elem.attrib.get(
                            "refUri",
                            ""
                        )

                        app_ref = (
                            extract_application_from_refuri(
                                refuri
                            )
                        )

                        if (
                            app_ref
                            not in app_map
                        ):

                            continue

                        app = app_map[
                            app_ref
                        ]

                        folder = (
                            find_application_folder(
                                extracted_iar,
                                app_ref
                            )
                        )

                        dbaas = (
                            analyze_dbaas_jca(
                                folder
                            )
                        )

                        desc = build_action_description(
                            app,
                            dbaas
                        )

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                app["Invoke"],

                            "Descripción de la Acción":
                                desc
                        })

                        seq += 1

                    # =====================================
                    # ASSIGNMENT
                    # =====================================

                    elif tag in [

                        "assign",

                        "assignment"
                    ]:

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "Assignment",

                            "Descripción de la Acción":
                                "Se realiza un Assignment."
                        })

                        seq += 1

                    # =====================================
                    # SWITCH
                    # =====================================

                    elif tag in [

                        "switch",

                        "router"
                    ]:

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "Switch",

                            "Descripción de la Acción":
                                "Se realiza un Switch."
                        })

                        seq += 1

                    # =====================================
                    # STAGE FILE
                    # =====================================

                    elif tag == "stageFile":

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                "stageFile",

                            "Descripción de la Acción":
                                "Se realiza un stageFile."
                        })

                        seq += 1

                    # =====================================
                    # FOR
                    # =====================================

                    elif tag == "for":

                        refuri = elem.attrib.get(
                            "refUri",
                            ""
                        )

                        processor_name = ""
                        iteration_name = ""
                        parallel = "false"

                        for proc in flow.iter():

                            try:

                                proc_tag = clean_tag(
                                    proc.tag
                                )

                            except:

                                continue

                            if proc_tag != "processor":

                                continue

                            if proc.attrib.get(
                                "name",
                                ""
                            ) != refuri:

                                continue

                            for child in proc.iter():

                                try:

                                    child_tag = clean_tag(
                                        child.tag
                                    )

                                except:

                                    continue

                                if child_tag.lower() == "processorname":

                                    if child.text:

                                        processor_name = (
                                            child.text.strip()
                                        )

                                elif child_tag.lower() == "subrole":

                                    if child.text:

                                        iteration_name = (
                                            child.text.strip()
                                        )

                                elif child_tag.lower() == "property":

                                    if (
                                        child.attrib.get(
                                            "name",
                                            ""
                                        )
                                        ==
                                        "parallel"
                                    ):

                                        parallel = child.attrib.get(
                                            "value",
                                            "false"
                                        )

                            break

                        mode = (
                            "paralela"
                            if parallel.lower() == "true"
                            else "secuencial"
                        )

                        if iteration_name:

                            description = (

                                f"Ahora comienza un For "
                                f"que itera sobre "
                                f"{iteration_name} "
                                f"de manera {mode}."
                            )

                        else:

                            description = (
                                "Se realiza un For."
                            )

                        rows.append({

                            "Nro. Secuencia":
                                seq,

                            "Nombre Acción":
                                processor_name or "For",

                            "Descripción de la Acción":
                                description
                        })

                        seq += 1

                # =========================================
                # TABLE
                # =========================================
                print("========== ROWS ==========")

                for row in rows:
                    print(row)

                print("==========================")

                add_process_table(
                    document,
                    rows
                )

                # =========================================
                # DIAGRAM
                # =========================================

                document.add_paragraph("")

                document.add_paragraph(
                    "Diagrama de Secuencia del Endpoint:"
                )

                integration_type = "REST"

                if is_scheduled_integration(
                    extracted_iar
                ):

                    integration_type = "scheduled"

                img_path = (
                    generate_sequence_diagram_png(

                        endpoint,

                        rows,

                        integration_type
                    )
                )

                if img_path and os.path.exists(
                    img_path
                ):

                    add_framed_image(

                        document,

                        img_path,

                        width=Inches(6.5)
                    )

                    # =================================
                    # CLEAN TEMP IMAGE
                    # =================================

                    try:

                        image_dir = os.path.dirname(
                            image_path
                        )

                        if os.path.exists(
                            image_path
                        ):
                            os.remove(
                                image_path
                            )

                        import shutil

                        shutil.rmtree(
                            image_dir,
                            ignore_errors=True
                        )

                    except Exception as cleanup_error:

                        print(
                            "[VB CLEANUP]",
                            cleanup_error
                        )


                # =========================================
                # FREQUENCY
                # =========================================

                if is_scheduled_integration(
                    extracted_iar
                ):

                    schedule_info = (
                        get_schedule_information(
                            extracted_iar
                        )
                    )

                    frequency = (
                        schedule_info.get(
                            "ical_expression",
                            "No definida"
                        )
                    )

                else:

                    frequency = "No aplica"

                document.add_paragraph(

                    f"Frecuencia de Ejecución\t: "
                    f"{frequency}"
                )

                document.add_paragraph("")
                document.add_paragraph("")
               
    # =====================================================
    # 3 DISEÑO DE PANTALLA
    # =====================================================
    document.add_page_break()
    create_header(
        document,
        "3\tDiseño de Pantalla"
    )

    screen_counter = 1

    # =====================================================
    # VALIDATION
    # =====================================================

    if not visual_builder_apps:

        document.add_paragraph(
            "No se encontraron aplicaciones Visual Builder."
        )

    else:

        # =================================================
        # NORMALIZE
        # =================================================

        if not isinstance(
            visual_builder_apps,
            list
        ):

            visual_builder_apps = [
                visual_builder_apps
            ]

        # =================================================
        # APPLICATIONS
        # =================================================

        for vb_zip in visual_builder_apps:

            try:

                # =========================================
                # EXTRACT
                # =========================================

                extraction_metadata = (
                    build_extraction_metadata(
                        vb_zip
                    )
                )

                # =========================================
                # PARSE PAGES
                # =========================================

                virtual_apps = (
                    build_virtual_apps(
                        extraction_metadata
                    )
                )

                for vb_metadata in virtual_apps:

                    shell_html = vb_metadata.get(
                        "shell_html",
                        ""
                    )

                    app_css = vb_metadata.get(
                        "app_css",
                        ""
                    )

                    resources_path = vb_metadata.get(
                        "resources_path"
                    )

                    pages = vb_metadata.get(
                        "pages",
                        []
                    )

                    page = pages[0]

                    page_name = page.get(
                        "name",
                        ""
                    )

                    page_html = page.get(
                        "html",
                        ""
                    )

                    document.add_paragraph(
                        page_name
                    )

                    complete_html = build_complete_html(
                        shell_html=shell_html,
                        page_html=page_html,
                        app_css=app_css
                    )

                    image_path = render_html_to_image(
                        html_content=complete_html,
                        resources_path=resources_path
                    )

                total_pages = len(
                    pages
                )

                page_index = 0

                current_step = 0

                steps_per_page = 5

                total_steps = (
                    total_pages
                    * steps_per_page
                )


                total_pages = len(
                    pages
                )

                page_index = 0

                current_step = 0

                steps_per_page = 5

                total_steps = (
                    total_pages
                    * steps_per_page
                )

                # =========================================
                # NO PAGES
                # =========================================

                if not pages:

                    document.add_paragraph(
                        "No se encontraron páginas Visual Builder."
                    )

                    continue

                # =========================================
                # LOOP PAGES
                # =========================================

                for page in pages:

                    page_index += 1

                    page_name = page.get(
                        "name",
                        f"Pantalla {screen_counter}"
                    )

                    update_activity(
                        job_id,
                        (
                            f"[{page_index}/{total_pages}] "
                            f"Iniciando pantalla: {page_name}"
                        ),
                        page_index,
                        total_pages
                    )

                    page_html = page.get(
                        "html",
                        ""
                    )

                    if screen_counter == 1:

                        if hasattr(
                            vb_zip,
                            "name"
                        ):

                            app_name = vb_zip.name

                        else:

                            app_name = os.path.basename(
                                str(vb_zip)
                            )

                        app_name = app_name.replace(
                            ".zip",
                            ""
                        )

                        create_header(
                            document,
                            f"3.1\t{app_name}",
                            size=14
                        )

                    document.add_paragraph(
                        page_name
                    )

                    try:

                        # =================================
                        # BUILD HTML
                        # =================================

                        update_activity(
                            job_id,
                            f"[{page_index}/{total_pages}] Construyendo HTML: {page_name}"
                        )

                        current_step += 1

                        advance_step(
                            job_id,
                            int(
                                (current_step / total_steps)
                                * 100
                            )
                        )

                        complete_html = (
                            build_complete_html(

                                shell_html=
                                    shell_html,

                                page_html=
                                    page_html,

                                app_css=
                                    app_css
                            )
                        )

                        # =================================
                        # RENDER IMAGE
                        # =================================

                        try:

                            update_activity(
                                job_id,
                                f"[{page_index}/{total_pages}] Generando imagen: {page_name}"
                            )

                            current_step += 1

                            advance_step(
                                job_id,
                                int(
                                    (current_step / total_steps)
                                    * 100
                                )
                            )

                            image_path = render_html_to_image(
                                html_content=complete_html,
                                resources_path=resources_path
                            )

                        except Exception as e:

                            image_path = None

                            document.add_paragraph(
                                f"No fue posible generar la captura Visual Builder: {str(e)}"
                            )

                        # =================================
                        # INSERT IMAGE
                        # =================================

                        if (

                            image_path

                            and

                            os.path.exists(
                                image_path
                            )
                        ):

                            update_activity(
                                job_id,
                                f"[{page_index}/{total_pages}] Insertando captura: {page_name}"
                            )

                            current_step += 1

                            advance_step(
                                job_id,
                                int(
                                    (current_step / total_steps)
                                    * 100
                                )
                            )

                            add_framed_image(

                                document,

                                image_path,

                                width=Inches(6.5)
                            )


                            # =============================
                            # BUTTONS
                            # =============================

                            buttons = page.get(
                                "buttons",
                                []
                            )

                            if buttons:

                                current_step += 1

                                advance_step(
                                    job_id,
                                    int(
                                        (current_step / total_steps)
                                        * 100
                                    )
                                )

                                document.add_paragraph(
                                    "• Lista de Botones"
                                )

                                button_table = (
                                    document.add_table(
                                        rows=1,
                                        cols=2
                                    )
                                )

                                document.add_paragraph("")

                                button_table.style = (
                                    "Table Grid"
                                )

                                hdr = (
                                    button_table.rows[0].cells
                                )

                                hdr[0].text = "BOTÓN"
                                hdr[1].text = "DESCRIPCIÓN"

                                apply_header_style(
                                    hdr[0]
                                )

                                apply_header_style(
                                    hdr[1]
                                )

                                for btn in buttons:

                                    row = (
                                        button_table.add_row().cells
                                    )

                                    row[0].text = (
                                        btn.get(
                                            "name",
                                            ""
                                        )
                                    )

                                    row[1].text = (
                                        btn.get(
                                            "description",
                                            ""
                                        )
                                    )

                            # =============================
                            # PARAMETERS
                            # =============================
                            
                            current_step += 1

                            advance_step(
                                job_id,
                                int(
                                    (current_step / total_steps)
                                    * 100
                                )
                            )

                            document.add_paragraph(
                                "• Lista de Parámetros"
                            )

                            parameter_table = (
                                document.add_table(
                                    rows=2,
                                    cols=2
                                )
                            )

                            document.add_paragraph("")

                            parameter_table.style = (
                                "Table Grid"
                            )

                            hdr = (
                                parameter_table.rows[0].cells
                            )

                            hdr[0].text = "PARÁMETROS"
                            hdr[1].text = "DESCRIPCIÓN"

                            apply_header_style(
                                hdr[0]
                            )

                            apply_header_style(
                                hdr[1]
                            )

                            row = (
                                parameter_table.rows[1].cells
                            )

                            row[0].text = (
                                "Todos los campos en la búsqueda"
                            )

                            row[1].text = (
                                "Todos los campos en la búsqueda"
                            )

                            # =============================
                            # FIELDS
                            # =============================

                            current_step += 1

                            advance_step(
                                job_id,
                                int(
                                    (current_step / total_steps)
                                    * 100
                                )
                            )

                            fields = page.get(
                                "table_columns",
                                []
                            )

                            if fields:

                                document.add_paragraph(
                                    "• Lista de Campos"
                                )

                                field_table = (
                                    document.add_table(
                                        rows=1,
                                        cols=2
                                    )
                                )

                                document.add_paragraph("")

                                field_table.style = (
                                    "Table Grid"
                                )

                                hdr = (
                                    field_table.rows[0].cells
                                )

                                hdr[0].text = "CAMPOS"
                                hdr[1].text = "DESCRIPCIÓN"

                                apply_header_style(
                                    hdr[0]
                                )

                                apply_header_style(
                                    hdr[1]
                                )

                                for field in fields:

                                    row = (
                                        field_table.add_row().cells
                                    )

                                    row[0].text = (
                                        field.get(
                                            "name",
                                            ""
                                        )
                                    )

                                    row[1].text = (
                                        field.get(
                                            "description",
                                            ""
                                        )
                                    )

                                update_activity(
                                    job_id,
                                    f"[{page_index}/{total_pages}] Pantalla completada: {page_name}"
                                )

                        else:

                            document.add_paragraph(
                                "No fue posible generar la imagen."
                            )

                    except Exception as e:

                        document.add_paragraph(
                            f"Error renderizando pantalla: {str(e)}"
                        )

                    document.add_paragraph("")
                    document.add_page_break()

                    screen_counter += 1

                    current_step += 1

                    advance_step(
                        job_id,
                        int(
                            (current_step / total_steps)
                            * 100
                        )
                    )

            except Exception as e:

                document.add_paragraph(
                    f"Error procesando Visual Builder: {str(e)}"
                )

            try:

                extraction_root = extraction_metadata.get(
                    "root_path"
                )

                if extraction_root:

                    import shutil

                    shutil.rmtree(
                        extraction_root,
                        ignore_errors=True
                    )

            except Exception as cleanup_error:

                print(
                    "[VB EXTRACT CLEANUP]",
                    cleanup_error
                )


    if job_id and visual_builder_apps:

        advance_step(

            job_id,

            "Visual Builder Procesado"
        )

    # =====================================================
    # 4 DISEÑO DE SERVICIOS
    # =====================================================
    document.add_page_break()
    create_header(
        document,
        "4\tDiseño de Interface"
    )

    add_description_box(

        document,

        "En esta sección se describe "
        "todo aquello que se utilice "
        "en el desarrollo como un "
        "servicio de interfaz "
        "(Ejem: SOAP, REST)."
    )

    document.add_paragraph("")

    create_header(
        document,
        "4.1\tDiseño de Servicio",
        size=14
    )

    if not package_path:

        document.add_paragraph(
            "No aplica."
        )

    else:

        if not integrations_data:

            document.add_paragraph(
                "No se encontraron integraciones."
            )

        else:

            for integration in integrations_data:

                if integration.get(
                    "scheduled",
                    False
                ):

                    continue
                try:

                    add_service_design_section(

                        document,

                        integration.get(
                            "path"
                        ),

                        integration.get(
                            "name"
                        ),

                        integration.get(
                            "version",
                            "01.00.0000"
                        )
                    )

                    document.add_paragraph("")
                    document.add_paragraph("")

                except Exception as e:

                    document.add_paragraph(

                        f"Error generando diseño "
                        f"de servicio: {str(e)}"
                    )

    if job_id and package_path:

        advance_step(

            job_id,

            "Diseño de Interfaces Procesado"
        )

    # =====================================================
    # 5 DISEÑO DEL REPORTE
    # =====================================================

    bip_metadata = {}

    if bip_files:

        try:
            document.add_page_break()
            # =============================================
            # BUILD ARTIFACT TREE
            # =============================================

            artifact_tree = build_bip_artifact_tree(
                bip_files
            )

            # =============================================
            # BUILD METADATA
            # =============================================

            bip_metadata = build_bip_metadata(
                artifact_tree
            )

            # =============================================
            # REPORT SECTION
            # =============================================

            add_report_design_section(
                document,
                bip_metadata
            )

        except Exception as e:

            document.add_page_break()

            create_header(
                document,
                "5\tDiseño del Reporte"
            )

            document.add_paragraph(
                f"Error procesando reportes BI Publisher: "
                f"{str(e)}"
            )

    # =====================================================
    # 6 SENTENCIAS SQL
    # =====================================================

    if bip_files and bip_metadata:

        try:
            document.add_page_break()
            add_sql_design_section(
                document,
                bip_metadata
            )

        except Exception as e:
            document.add_page_break()
            create_header(
                document,
                "6\tSentencias SQL"
            )

            document.add_paragraph(
                f"Error generando "
                f"sentencias SQL: "
                f"{str(e)}"
            )

    # =====================================================
    # 7 DISEÑO DE BASE DE DATOS
    # =====================================================

    if database_metadata:

        try:
            document.add_page_break()
            add_database_design_section(

                document,

                database_metadata,

                database_export_info
            )

        except Exception as e:
            document.add_page_break()
            create_header(
                document,
                "7\tDiseño de Base de Datos"
            )

            document.add_paragraph(

                f"Error generando "
                f"diseño de base de datos: "
                f"{str(e)}"
            )

    # =====================================================
    # 8 CONTROL DE ACCESOS
    # =====================================================

    document.add_page_break()

    create_header(
        document,
        "8\tControl de Accesos"
    )

    add_description_box(

        document,

        "En esta sección se lista los usuarios "
        "y tipos de roles que deben ser considerados "
        "en los diferentes componentes para el "
        "funcionamiento del desarrollo."
    )

    document.add_paragraph("")

    access_table = document.add_table(
        rows=1,
        cols=4
    )

    access_table.style = "Table Grid"

    hdr = access_table.rows[0].cells

    hdr[0].text = "Componente"
    hdr[1].text = "Nombre de Usuario"
    hdr[2].text = "Método de Control de Acceso"
    hdr[3].text = "Comentarios"

    for cell in hdr:

        apply_header_style(
            cell
        )

    access_rows = []

    access_rows.append({

        "component":
            "ERP Cloud",

        "user":
            "<Usuario de ERP>",

        "method":
            "Permite acceder al sistema para ejecutar los reportes",

        "comment":
            "El usuario debe contar con privilegios para acceder al OTBI, modificar, crear y eliminar reportes."
    })

    if (

        "OIC" in selected_components

        or

        "Visual Builder" in selected_components

    ):

        access_rows.append({

            "component":
                "OIC",

            "user":
                "<Usuario de OIC>",

            "method":
                "Permite acceder al sistema para ejecutar las integraciones",

            "comment":
                "El usuario debe contar con privilegios para acceder o modificar las integraciones."
        })

    if "Objetos BD" in selected_components:

        access_rows.append({

            "component":
                "BD",

            "user":
                "<Usuario de BD>",

            "method":
                "Permite acceder a los objetos de base de datos según el esquema establecido",

            "comment":
                "El usuario debe contar con privilegios para acceder, crear, modificar y eliminar tablas, secuencias, vistas, paquetes, etc."
        })

    if "APEX" in selected_components:

        access_rows.append({

            "component":
                "APEX",

            "user":
                "<Usuario de APEX>",

            "method":
                "Permite acceder a los objetos de base de datos según el esquema establecido además de las aplicaciones creadas en el Workspace del usuario",

            "comment":
                "El usuario debe contar con privilegios para acceder, crear, modificar y eliminar tablas, secuencias, vistas, paquetes, etc."
        })

    for item in access_rows:

        row = access_table.add_row().cells

        row[0].text = item["component"]
        row[1].text = item["user"]
        row[2].text = item["method"]
        row[3].text = item["comment"]

    # =====================================================
    # 10 TEMAS ABIERTOS Y CERRADOS
    # =====================================================

    document.add_page_break()

    create_header(
        document,
        "10\tTemas Abiertos y Cerrados"
    )

    # =====================================================
    # 10.1 TEMAS ABIERTOS
    # =====================================================

    create_header(
        document,
        "10.1\tTemas Abiertos"
    )

    open_table = document.add_table(
        rows=5,
        cols=6
    )

    open_table.style = "Table Grid"

    headers = [

        "ID",

        "Tema",

        "Resolución",

        "Responsabilidad",

        "Fecha Estimada",

        "Fecha Impactada"
    ]

    hdr = open_table.rows[0].cells

    for i in range(6):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )

    document.add_paragraph("")

    # =====================================================
    # 10.2 TEMAS CERRADOS
    # =====================================================

    create_header(
        document,
        "10.2\tTemas Cerrados"
    )

    closed_table = document.add_table(
        rows=5,
        cols=6
    )

    closed_table.style = "Table Grid"

    hdr = closed_table.rows[0].cells

    for i in range(6):

        hdr[i].text = headers[i]

        apply_header_style(
            hdr[i]
        )


    output = BytesIO()

    document.save(output)

    if job_id:

        advance_step(

            job_id,

            "Documento Word Generado"
        )

    output.seek(0)

    return output