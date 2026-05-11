# =========================================================
# word_utils.py
# =========================================================

import os
import re

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn

from xml_utils import (
    find_all_iar_files,
    extract_iar,
    get_project_name,
    read_project_xml,
    build_application_map,
    get_project_root,
    clean_tag,
    extract_application_from_refuri,
    find_application_folder,
    analyze_dbaas_jca,
    generate_description,
    get_integration_metadata,
    get_schedule_information,
    get_javascript_files,
    get_lookup_files,
    get_connection_xmls,
    get_request_response_schemas,
    parse_schema_elements,
    get_active_integrations_only
)

from mermaid_utils import (
    generate_sequence_diagram_png
)


# =========================================================
# ADD TITLE
# =========================================================

def add_title(
    document,
    text,
    size=15
):

    p = document.add_paragraph()

    run = p.add_run(text)

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)

    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )


# =========================================================
# ADD HEADING
# =========================================================

def add_heading(
    document,
    text,
    level=1
):

    p = document.add_paragraph()

    run = p.add_run(text)

    run.bold = True
    run.font.name = "Arial"

    if level == 1:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


# =========================================================
# CREATE TABLE
# =========================================================

def create_table(
    document,
    headers,
    rows
):

    table = document.add_table(
        rows=1,
        cols=len(headers)
    )

    table.style = 'Table Grid'

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    hdr = table.rows[0].cells

    for i in range(len(headers)):

        hdr[i].text = headers[i]

        for paragraph in hdr[i].paragraphs:

            paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            if paragraph.runs:

                paragraph.runs[0].bold = True

        shading_elm = parse_xml(
            r'<w:shd {} w:fill="D9D9D9"/>'.format(
                nsdecls('w')
            )
        )

        hdr[i]._tc.get_or_add_tcPr().append(
            shading_elm
        )

    for row in rows:

        cells = table.add_row().cells

        for idx, value in enumerate(row):

            cells[idx].text = str(value)


# =========================================================
# BUILD FLOW ROWS
# =========================================================

def build_flow_rows(
    flow,
    app_map,
    extracted_iar
):

    rows = []

    seq = 1

    for elem in flow.iter():

        if clean_tag(
            elem.tag
        ) != "invoke":
            continue

        refuri = elem.attrib.get(
            "refUri",
            ""
        )

        app_ref = (
            extract_application_from_refuri(
                refuri
            )
        )

        if app_ref not in app_map:
            continue

        app = app_map[
            app_ref
        ]

        folder = (
            find_application_folder(
                extracted_iar,
                app_ref
            )
        )

        dbaas = (
            analyze_dbaas_jca(
                folder
            )
        )

        desc = (
            f"De tipo "
            f"{app['Tipo']}"
        )

        if dbaas["Operacion"]:

            desc += (
                f", realiza "
                f"operación "
                f"{dbaas['Operacion']}"
            )

        if dbaas["Tabla"]:

            desc += (
                f" a la tabla "
                f"{dbaas['Tabla']}"
            )

        if dbaas["SQL"]:

            desc += (
                f", query "
                f"'{dbaas['SQL']}'"
            )

        if (
            dbaas["Package"]
            and
            dbaas["Procedure"]
        ):

            desc += (
                f", package "
                f"{dbaas['Package']}."
                f"{dbaas['Procedure']}"
            )

        desc += "."

        rows.append({

            "Nro. Secuencia": seq,

            "Nombre Acción":
                app["Invoke"],

            "Descripción de la Acción":
                desc
        })

        seq += 1

    return rows


# =========================================================
# ADD CONNECTION SECTION
# =========================================================

def add_connections_section(
    document,
    connection_files
):

    add_heading(
        document,
        "Conexiones Utilizadas",
        2
    )

    if not connection_files:

        document.add_paragraph(
            "La integración no utiliza conexiones."
        )

        return

    rows = []

    for conn in connection_files:

        rows.append([

            conn.get(
                "name",
                ""
            ),

            conn.get(
                "type",
                ""
            ),

            conn.get(
                "security_policy",
                ""
            ),

            conn.get(
                "agent_group",
                ""
            )
        ])

    create_table(

        document,

        [
            "Conexión",
            "Tipo",
            "Seguridad",
            "Agent Group"
        ],

        rows
    )


# =========================================================
# ADD JS SECTION
# =========================================================

def add_javascript_section(
    document,
    js_files
):

    add_heading(
        document,
        "Uso de Javascript",
        2
    )

    if not js_files:

        document.add_paragraph(
            "La integración no utiliza Javascript."
        )

        return

    document.add_paragraph(
        "La integración utiliza los siguientes archivos Javascript:"
    )

    for js in js_files:

        document.add_paragraph(
            f"• {js}"
        )


# =========================================================
# ADD LOOKUP SECTION
# =========================================================

def add_lookup_section(
    document,
    lookup_files
):

    add_heading(
        document,
        "Lookups Utilizados",
        2
    )

    if not lookup_files:

        document.add_paragraph(
            "La integración no utiliza Lookups."
        )

        return

    for lookup in lookup_files:

        document.add_paragraph(
            f"• {lookup}"
        )


# =========================================================
# ADD DESIGN SERVICE SECTION
# =========================================================

def add_design_service_section(
    document,
    integration_name,
    version,
    schemas
):

    add_heading(
        document,
        "4.1 Diseño de Servicio",
        1
    )

    document.add_paragraph(
        f"• {integration_name} | {version}"
    )

    if not schemas:

        document.add_paragraph(
            "No se encontraron request/response."
        )

        return

    for schema in schemas:

        endpoint_name = schema.get(
            "endpoint",
            ""
        )

        has_request = bool(
            schema.get(
                "request_elements"
            )
        )

        has_response = bool(
            schema.get(
                "response_elements"
            )
        )

        req_res = []

        if has_request:
            req_res.append(
                "Request"
            )

        if has_response:
            req_res.append(
                "Response"
            )

        req_res_text = "/".join(
            req_res
        )

        document.add_paragraph(
            f"{endpoint_name} - {req_res_text}"
        )

        document.add_paragraph(
            "Tipo de Dato : JSON"
        )

        rows = []

        if has_request:

            request_data = "\n".join([

                f"{x['name']} ({x['type']})"

                for x in schema[
                    "request_elements"
                ]
            ])

            rows.append([
                "Request",
                request_data
            ])

        if has_response:

            response_data = "\n".join([

                f"{x['name']} ({x['type']})"

                for x in schema[
                    "response_elements"
                ]
            ])

            rows.append([
                "Response",
                response_data
            ])

        create_table(

            document,

            [
                "Tipo",
                "Wrapper"
            ],

            rows
        )


# =========================================================
# GENERATE WORD DOCUMENT
# =========================================================

def generate_word_document(
    package_path
):

    document = Document()

    style = document.styles['Normal']

    style.font.name = 'Arial'

    style._element.rPr.rFonts.set(
        qn('w:eastAsia'),
        'Arial'
    )

    style.font.size = Pt(10)

    # =====================================================
    # TITLE
    # =====================================================

    add_title(

        document,

        "2.5\tDiagrama de bloques "
        "de construcción"
    )

    # =====================================================
    # FIND IARS
    # =====================================================

    iar_files = find_all_iar_files(
        package_path
    )

    # =====================================================
    # FILTER ACTIVE
    # =====================================================

    active_integrations = (
        get_active_integrations_only(
            iar_files
        )
    )

    # =====================================================
    # PROCESS ACTIVE IARS
    # =====================================================

    for iar in active_integrations:

        extracted_iar = extract_iar(iar)

        integration_name = (
            get_project_name(
                extracted_iar
            )
        )

        metadata = (
            get_integration_metadata(
                extracted_iar
            )
        )

        integration_type = metadata.get(
            "integration_type",
            "REQUEST"
        )

        version = metadata.get(
            "version",
            ""
        )

        # =================================================
        # TITLE
        # =================================================

        add_heading(
            document,
            f"• {integration_name}",
            1
        )

        document.add_paragraph(
            f"Versión: {version}"
        )

        document.add_paragraph(
            f"Tipo: {integration_type}"
        )

        applications = (
            read_project_xml(
                extracted_iar
            )
        )

        app_map = build_application_map(
            applications
        )

        root = get_project_root(
            extracted_iar
        )

        endpoint_flows = {}

        # =================================================
        # APP DRIVEN
        # =================================================

        for orchestration in root.iter():

            if clean_tag(
                orchestration.tag
            ) != "orchestration":
                continue

            for pick in orchestration:

                if clean_tag(
                    pick.tag
                ) != "pick":
                    continue

                for pick_receive in pick:

                    if clean_tag(
                        pick_receive.tag
                    ) != "pickReceive":
                        continue

                    refuri = (
                        pick_receive.attrib.get(
                            "refUri",
                            ""
                        )
                    )

                    endpoint_app = (
                        extract_application_from_refuri(
                            refuri
                        )
                    )

                    if (
                        endpoint_app
                        not in app_map
                    ):
                        continue

                    endpoint_name = (
                        app_map[
                            endpoint_app
                        ]["Invoke"]
                    )

                    endpoint_flows[
                        endpoint_name
                    ] = pick_receive

        # =================================================
        # SCHEDULED
        # =================================================

        if not endpoint_flows:

            endpoint_flows[
                integration_name
            ] = root

        # =================================================
        # FLOW PROCESS
        # =================================================

        for endpoint, flow in (
            endpoint_flows.items()
        ):

            description = (
                generate_description(
                    endpoint,
                    flow,
                    applications,
                    extracted_iar
                )
            )

            document.add_paragraph(
                description
            )

            document.add_paragraph(
                "A continuación, "
                "se describe cada "
                "una de las acciones "
                "del flujo de integración:"
            )

            rows = []

            seq = 1

            rows.append({

                "Nro. Secuencia": seq,

                "Nombre Acción":
                    endpoint,

                "Descripción de la Acción":
                    (
                        "Integración programada."
                        if integration_type.upper() == "SCHEDULED"
                        else "Endpoint REST."
                    )
            })

            seq += 1

            flow_rows = build_flow_rows(
                flow,
                app_map,
                extracted_iar
            )

            for row in flow_rows:

                row[
                    "Nro. Secuencia"
                ] = seq

                rows.append(row)

                seq += 1

            create_table(

                document,

                [
                    "Nro. Secuencia",
                    "Nombre Acción",
                    "Descripción de la Acción"
                ],

                [

                    [
                        x[
                            "Nro. Secuencia"
                        ],

                        x[
                            "Nombre Acción"
                        ],

                        x[
                            "Descripción de la Acción"
                        ]
                    ]

                    for x in rows
                ]
            )

            # =============================================
            # DIAGRAM
            # =============================================

            document.add_paragraph(
                "Diagrama de Secuencia:"
            )

            img_path = (
                generate_sequence_diagram_png(
                    endpoint,
                    rows,
                    integration_type
                )
            )

            document.add_picture(
                img_path,
                width=Inches(6.5)
            )

            # =============================================
            # FREQUENCY
            # =============================================

            schedule = (
                get_schedule_information(
                    extracted_iar
                )
            )

            if (
                integration_type.upper()
                == "SCHEDULED"
            ):

                document.add_paragraph(
                    f"Frecuencia de Ejecución : "
                    f"{schedule}"
                )

            else:

                document.add_paragraph(
                    "Frecuencia de Ejecución : "
                    "Por demanda"
                )

            document.add_paragraph("")

        # =================================================
        # JAVASCRIPT
        # =================================================

        js_files = get_javascript_files(
            extracted_iar
        )

        add_javascript_section(
            document,
            js_files
        )

        # =================================================
        # LOOKUPS
        # =================================================

        lookup_files = get_lookup_files(
            extracted_iar
        )

        add_lookup_section(
            document,
            lookup_files
        )

        # =================================================
        # CONNECTIONS
        # =================================================

        connection_files = (
            get_connection_xmls(
                extracted_iar
            )
        )

        add_connections_section(
            document,
            connection_files
        )

        # =================================================
        # SERVICE DESIGN
        # =================================================

        schemas = (
            get_request_response_schemas(
                extracted_iar
            )
        )

        add_design_service_section(

            document,

            integration_name,

            version,

            schemas
        )

        document.add_page_break()

    # =====================================================
    # OUTPUT
    # =====================================================

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output