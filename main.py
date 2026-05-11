# =========================================================
# REQUERIMIENTOS
# =========================================================
# pip install streamlit pandas python-docx
#
# Además instalar:
# npm install -g @mermaid-js/mermaid-cli
#
# Verificar:
# mmdc -h

# =========================================================
# APP.PY
# =========================================================

import streamlit as st
import zipfile
import tempfile
import os
import re
import subprocess
import shutil
import xml.etree.ElementTree as ET

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn


# =========================================================
# CONFIG
# =========================================================

st.set_page_config(
    page_title="Oracle PAR Documentation Generator",
    layout="wide"
)

st.title("📘 Oracle PAR Documentation Generator")


# =========================================================
# EXTRACT PACKAGE
# =========================================================

def extract_package(uploaded_file):

    temp_dir = tempfile.mkdtemp()

    with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    return temp_dir


# =========================================================
# FIND ALL IAR FILES
# =========================================================

def find_all_iar_files(base_path):

    iar_files = []

    for root, dirs, files in os.walk(base_path):

        for file in files:

            if file.lower().endswith(".iar"):

                iar_files.append(
                    os.path.join(root, file)
                )

    return iar_files


# =========================================================
# EXTRACT IAR
# =========================================================

def extract_iar(iar_path):

    temp_dir = tempfile.mkdtemp()

    with zipfile.ZipFile(iar_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    return temp_dir


# =========================================================
# CREATE ZIP FROM FOLDER
# =========================================================

def create_zip_from_folder(folder_path):

    memory_file = BytesIO()

    with zipfile.ZipFile(
        memory_file,
        'w',
        zipfile.ZIP_DEFLATED
    ) as zipf:

        for root, dirs, files in os.walk(folder_path):

            for file in files:

                full_path = os.path.join(
                    root,
                    file
                )

                relative_path = os.path.relpath(
                    full_path,
                    folder_path
                )

                zipf.write(
                    full_path,
                    relative_path
                )

    memory_file.seek(0)

    return memory_file


# =========================================================
# FIND FILE
# =========================================================

def find_file(base_path, filename):

    for root, dirs, files in os.walk(base_path):

        for file in files:

            if file.lower() == filename.lower():

                return os.path.join(root, file)

    return None


# =========================================================
# CLEAN XML TAG
# =========================================================

def clean_tag(tag):

    if "}" in tag:
        return tag.split("}", 1)[1]

    return tag


# =========================================================
# EXTRACT APP FROM REFURI
# =========================================================

def extract_application_from_refuri(refuri):

    if not refuri:
        return None

    match = re.search(
        r'(application_\d+)',
        refuri
    )

    if match:
        return match.group(1)

    return None


# =========================================================
# CAMEL TO SNAKE
# =========================================================

def camel_to_snake_upper(name):

    s1 = re.sub(
        '(.)([A-Z][a-z]+)',
        r'\1_\2',
        name
    )

    s2 = re.sub(
        '([a-z0-9])([A-Z])',
        r'\1_\2',
        s1
    )

    return s2.upper()


# =========================================================
# FIND APPLICATION FOLDER
# =========================================================

def find_application_folder(
    base_path,
    application_name
):

    for root, dirs, files in os.walk(base_path):

        for d in dirs:

            if d.lower() == application_name.lower():

                return os.path.join(root, d)

    return None


# =========================================================
# FIND JCA
# =========================================================

def find_jca_file(application_folder):

    if not application_folder:
        return None

    for root, dirs, files in os.walk(application_folder):

        for file in files:

            if file.endswith(".jca"):

                return os.path.join(
                    root,
                    file
                )

    return None


# =========================================================
# DETECT CONNECTION TYPE
# =========================================================

def detect_connection_type(
    adapter_code,
    adapter_type
):

    value = (
        f"{adapter_code} {adapter_type}"
    ).lower()

    if (
        "dbaas" in value
        or "db" in value
    ):
        return "DBaaS"

    if "rest" in value:
        return "REST"

    if "soap" in value:
        return "SOAP"

    if "ftp" in value:
        return "FTP"

    if "sftp" in value:
        return "SFTP"

    if "erp" in value:
        return "ERP Cloud"

    return "Unknown"


# =========================================================
# PROJECT ROOT
# =========================================================

def get_project_root(base_path):

    project_file = find_file(
        base_path,
        "project.xml"
    )

    if not project_file:
        return None

    tree = ET.parse(project_file)

    return tree.getroot()


# =========================================================
# PROJECT NAME
# =========================================================

def get_project_name(base_path):

    project_file = find_file(
        base_path,
        "project.xml"
    )

    tree = ET.parse(project_file)

    root = tree.getroot()

    for elem in root.iter():

        if clean_tag(elem.tag) == "projectName":

            return elem.text

    return "UNKNOWN"


# =========================================================
# READ APPLICATIONS
# =========================================================

def read_project_xml(base_path):

    root = get_project_root(base_path)

    applications = []

    for elem in root.iter():

        if clean_tag(elem.tag) == "application":

            app_name = elem.attrib.get("name")

            invoke_name = ""
            adapter_type = ""
            adapter_code = ""

            for child in elem.iter():

                tag = clean_tag(child.tag)

                text = (
                    child.text.strip()
                    if child.text
                    else ""
                )

                if tag == "name":
                    invoke_name = text

                if tag == "type":
                    adapter_type = text

                if tag == "code":
                    adapter_code = text

            applications.append({

                "Application": app_name,

                "Invoke": invoke_name,

                "Tipo": detect_connection_type(
                    adapter_code,
                    adapter_type
                )
            })

    return applications


# =========================================================
# BUILD APP MAP
# =========================================================

def build_application_map(applications):

    return {
        app["Application"]: app
        for app in applications
    }


# =========================================================
# ANALYZE JCA
# =========================================================

def analyze_dbaas_jca(application_folder):

    result = {

        "Operacion": "",
        "Tabla": "",
        "Package": "",
        "Procedure": "",
        "SQL": ""
    }

    jca_file = find_jca_file(
        application_folder
    )

    if not jca_file:
        return result

    with open(
        jca_file,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as f:

        content = f.read()

    if (
        "DBStoredProcedureInteractionSpec"
        in content
    ):

        result["Operacion"] = (
            "Stored Procedure"
        )

        package_match = re.search(
            r'PackageName\" value=\"([^\"]+)\"',
            content
        )

        procedure_match = re.search(
            r'ProcedureName\" value=\"([^\"]+)\"',
            content
        )

        if package_match:
            result["Package"] = (
                package_match.group(1)
            )

        if procedure_match:
            result["Procedure"] = (
                procedure_match.group(1)
            )

    if (
        "DBPureSQLInteractionSpec"
        in content
    ):

        result["Operacion"] = (
            "SQL Puro"
        )

        sql_match = re.search(
            r'SqlString\" value=\"([^\"]+)\"',
            content
        )

        if sql_match:

            result["SQL"] = (
                sql_match.group(1)
            )

            table_match = re.search(
                r'from\s+([a-zA-Z0-9_\.]+)',
                result["SQL"],
                re.IGNORECASE
            )

            if table_match:

                result["Tabla"] = (
                    table_match.group(1)
                )

    if (
        "DBReadInteractionSpec"
        in content
    ):

        if "QueryName" in content:

            result["Operacion"] = (
                "Select"
            )

        descriptor_match = re.search(
            r'DescriptorName\" value=\"([^\"]+)\"',
            content
        )

        if descriptor_match:

            descriptor = (
                descriptor_match.group(1)
            )

            if "." in descriptor:
                descriptor = descriptor.split(".")[-1]

            result["Tabla"] = (
                camel_to_snake_upper(
                    descriptor
                )
            )

    if (
        "DBWriteInteractionSpec"
        in content
    ):

        dml_match = re.search(
            r'DmlType\" value=\"([^\"]+)\"',
            content
        )

        if dml_match:

            result["Operacion"] = (
                dml_match.group(1)
                .capitalize()
            )

        descriptor_match = re.search(
            r'DescriptorName\" value=\"([^\"]+)\"',
            content
        )

        if descriptor_match:

            descriptor = (
                descriptor_match.group(1)
            )

            if "." in descriptor:
                descriptor = descriptor.split(".")[-1]

            result["Tabla"] = (
                camel_to_snake_upper(
                    descriptor
                )
            )

    return result


# =========================================================
# GENERATE DESCRIPTION
# =========================================================

def generate_description(
    endpoint_name,
    flow_node,
    applications,
    base_path
):

    app_map = build_application_map(
        applications
    )

    parts = []

    first_invoke = True

    def process_node(node):

        nonlocal first_invoke

        for child in list(node):

            tag = clean_tag(child.tag)

            if tag == "invoke":

                refuri = child.attrib.get(
                    "refUri",
                    ""
                )

                app_ref = (
                    extract_application_from_refuri(
                        refuri
                    )
                )

                if app_ref not in app_map:
                    continue

                app = app_map[app_ref]

                if first_invoke:

                    parts.append(
                        f"una conexión "
                        f"{app['Tipo']} "
                        f"llamada "
                        f"{app['Invoke']}"
                    )

                    first_invoke = False

                else:

                    parts.append(
                        f"se llama a una "
                        f"conexión "
                        f"{app['Tipo']} "
                        f"llamada "
                        f"{app['Invoke']}"
                    )

            elif tag == "assignment":

                parts.append(
                    "se realiza un Assignment"
                )

            elif tag == "router":

                parts.append(
                    "se realiza un Switch"
                )

            elif tag == "try":

                parts.append(
                    "se realiza un Scope"
                )

            process_node(child)

    process_node(flow_node)

    description = (
        f"El endpoint "
        f"{endpoint_name} "
        f"comienza con: "
    )

    description += ", ".join(parts)

    description += "."

    return description


# =========================================================
# SANITIZE MERMAID
# =========================================================

def sanitize_mermaid_text(text):

    if not text:
        return "UNKNOWN"

    text = re.sub(
        r'["\'<>]',
        '',
        text
    )

    text = text.replace("-", "_")
    text = text.replace(" ", "_")
    text = text.replace("/", "_")

    return text


# =========================================================
# GENERATE MERMAID PNG
# =========================================================

def generate_sequence_diagram_png(
    endpoint_name,
    rows
):

    mmdc_path = shutil.which(
        "mmdc"
    )

    if not mmdc_path:

        possible_paths = [

            r"C:\Users\LP-KQ-NEORA\AppData\Roaming\npm\mmdc.cmd",

            r"C:\Users\LP-KQ-NEORA\AppData\Roaming\npm\mmdc"

        ]

        for path in possible_paths:

            if os.path.exists(path):

                mmdc_path = path
                break

    if not mmdc_path:

        raise Exception(
            "No se encontró Mermaid CLI (mmdc)"
        )

    temp_dir = tempfile.mkdtemp()

    mmd_file = os.path.join(
        temp_dir,
        "diagram.mmd"
    )

    png_file = os.path.join(
        temp_dir,
        "diagram.png"
    )

    lines = []

    lines.append(
        "%%{init: {'theme':'neutral'}}%%"
    )

    lines.append(
        "sequenceDiagram"
    )

    lines.append(
        "autonumber"
    )

    endpoint_clean = sanitize_mermaid_text(
        endpoint_name
    )

    lines.append(
        "participant CLIENT"
    )

    lines.append(
        f"participant API as {endpoint_clean}"
    )

    participants = []

    for row in rows[1:]:

        desc = row[
            "Descripción de la Acción"
        ]

        tipo_match = re.search(
            r'De tipo ([^,\.]+)',
            desc
        )

        participant = "SYSTEM"

        if tipo_match:

            participant = sanitize_mermaid_text(
                tipo_match.group(1)
            )

        if participant not in participants:

            participants.append(
                participant
            )

    for participant in participants:

        lines.append(
            f"participant {participant}"
        )

    lines.append(
        "CLIENT->>API: Request"
    )

    for row in rows[1:]:

        action = sanitize_mermaid_text(
            row["Nombre Acción"]
        )

        desc = row[
            "Descripción de la Acción"
        ]

        tipo_match = re.search(
            r'De tipo ([^,\.]+)',
            desc
        )

        target = "SYSTEM"

        if tipo_match:

            target = sanitize_mermaid_text(
                tipo_match.group(1)
            )

        lines.append(
            f"API->>{target}: {action}"
        )

        lines.append(
            f"{target}-->>API: Response"
        )

    lines.append(
        "API-->>CLIENT: Response"
    )

    mermaid_code = "\n".join(
        lines
    )

    with open(
        mmd_file,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(
            mermaid_code
        )

    result = subprocess.run(

        [
            mmdc_path,
            "-i",
            mmd_file,
            "-o",
            png_file,
            "-t",
            "neutral",
            "-b",
            "white",
            "-s",
            "2"
        ],

        capture_output=True,
        text=True

    )

    if result.returncode != 0:

        st.error(
            result.stderr
        )

        raise Exception(
            result.stderr
        )

    if not os.path.exists(
        png_file
    ):

        raise Exception(
            "Mermaid no generó PNG"
        )

    return png_file


# =========================================================
# CREATE WORD
# =========================================================

def generate_word_document(
    package_path
):

    document = Document()

    style = document.styles['Normal']

    style.font.name = 'Arial'

    style._element.rPr.rFonts.set(
        qn('w:eastAsia'),
        'Arial'
    )

    style.font.size = Pt(10)

    p = document.add_paragraph()

    run = p.add_run(
        "2.5\tDiagrama de bloques "
        "de construcción"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)

    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    iar_files = find_all_iar_files(
        package_path
    )

    for iar in iar_files:

        extracted_iar = extract_iar(iar)

        integration_name = (
            get_project_name(
                extracted_iar
            )
        )

        p = document.add_paragraph()

        run = p.add_run(
            f"• {integration_name}"
        )

        run.bold = True

        applications = (
            read_project_xml(
                extracted_iar
            )
        )

        app_map = build_application_map(
            applications
        )

        root = get_project_root(
            extracted_iar
        )

        endpoint_flows = {}

        for orchestration in root.iter():

            if clean_tag(
                orchestration.tag
            ) != "orchestration":
                continue

            for pick in orchestration:

                if clean_tag(
                    pick.tag
                ) != "pick":
                    continue

                for pick_receive in pick:

                    if clean_tag(
                        pick_receive.tag
                    ) != "pickReceive":
                        continue

                    refuri = (
                        pick_receive.attrib.get(
                            "refUri",
                            ""
                        )
                    )

                    endpoint_app = (
                        extract_application_from_refuri(
                            refuri
                        )
                    )

                    if (
                        endpoint_app
                        not in app_map
                    ):
                        continue

                    endpoint_name = (
                        app_map[
                            endpoint_app
                        ]["Invoke"]
                    )

                    endpoint_flows[
                        endpoint_name
                    ] = pick_receive

        for endpoint, flow in (
            endpoint_flows.items()
        ):

            description = (
                generate_description(
                    endpoint,
                    flow,
                    applications,
                    extracted_iar
                )
            )

            document.add_paragraph(
                description
            )

            document.add_paragraph(
                "A continuación, "
                "se describe cada "
                "una de las acciones "
                "del flujo de integración:"
            )

            rows = []

            seq = 1

            rows.append({

                "Nro. Secuencia": seq,

                "Nombre Acción":
                    endpoint,

                "Descripción de la Acción":
                    "De tipo REST, endpoint."
            })

            seq += 1

            for elem in flow.iter():

                if clean_tag(
                    elem.tag
                ) != "invoke":
                    continue

                refuri = elem.attrib.get(
                    "refUri",
                    ""
                )

                app_ref = (
                    extract_application_from_refuri(
                        refuri
                    )
                )

                if (
                    app_ref
                    not in app_map
                ):
                    continue

                app = app_map[
                    app_ref
                ]

                folder = (
                    find_application_folder(
                        extracted_iar,
                        app_ref
                    )
                )

                dbaas = (
                    analyze_dbaas_jca(
                        folder
                    )
                )

                desc = (
                    f"De tipo "
                    f"{app['Tipo']}"
                )

                if dbaas["Operacion"]:

                    desc += (
                        f", realiza "
                        f"operación "
                        f"{dbaas['Operacion']}"
                    )

                if dbaas["Tabla"]:

                    desc += (
                        f" a la tabla "
                        f"{dbaas['Tabla']}"
                    )

                if dbaas["SQL"]:

                    desc += (
                        f" con el siguiente "
                        f"query "
                        f"'{dbaas['SQL']}'"
                    )

                if (
                    dbaas["Package"]
                    and
                    dbaas["Procedure"]
                ):

                    desc += (
                        f" en el paquete "
                        f"{dbaas['Package']}."
                        f"{dbaas['Procedure']}"
                    )

                desc += "."

                rows.append({

                    "Nro. Secuencia": seq,

                    "Nombre Acción":
                        app["Invoke"],

                    "Descripción de la Acción":
                        desc
                })

                seq += 1

            table = document.add_table(
                rows=1,
                cols=3
            )

            table.style = 'Table Grid'

            table.alignment = (
                WD_TABLE_ALIGNMENT.CENTER
            )

            hdr = table.rows[0].cells

            headers = [
                "Nro. Secuencia",
                "Nombre Acción",
                "Descripción de la Acción"
            ]

            for i in range(3):

                hdr[i].text = headers[i]

                for paragraph in hdr[i].paragraphs:

                    paragraph.alignment = (
                        WD_PARAGRAPH_ALIGNMENT.CENTER
                    )

                    run = paragraph.runs[0]

                    run.bold = True

                shading_elm = parse_xml(
                    r'<w:shd {} w:fill="D9D9D9"/>'.format(
                        nsdecls('w')
                    )
                )

                hdr[i]._tc.get_or_add_tcPr().append(
                    shading_elm
                )

            for row in rows:

                cells = table.add_row().cells

                cells[0].text = str(
                    row[
                        "Nro. Secuencia"
                    ]
                )

                cells[1].text = (
                    row[
                        "Nombre Acción"
                    ]
                )

                cells[2].text = (
                    row[
                        "Descripción de la Acción"
                    ]
                )

            document.add_paragraph(
                "Diagrama de "
                "Secuencia del Endpoint:"
            )

            img_path = (
                generate_sequence_diagram_png(
                    endpoint,
                    rows
                )
            )

            document.add_picture(
                img_path,
                width=Inches(6.5)
            )

            document.add_paragraph(
                "Frecuencia de "
                "Ejecución\t: "
                "Por demanda"
            )

            document.add_paragraph("")
            document.add_paragraph("")

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output


# =========================================================
# UI - PAR
# =========================================================

uploaded_file = st.file_uploader(
    "Subir archivo .par",
    type=["par"]
)

if uploaded_file:

    st.success(
        "Archivo PAR cargado correctamente."
    )

    # =====================================
    # DESCARGAR CONTENIDO DEL PAR
    # =====================================

    if st.button(
        "📦 Descargar contenido interno del PAR"
    ):

        try:

            package_path = extract_package(
                uploaded_file
            )

            zip_content = create_zip_from_folder(
                package_path
            )

            st.download_button(

                label="⬇️ Descargar contenido del PAR",

                data=zip_content,

                file_name="contenido_par.zip",

                mime="application/zip"
            )

        except Exception as e:

            st.error(
                f"Error extrayendo PAR: {str(e)}"
            )

    # =====================================
    # GENERAR WORD
    # =====================================

    if st.button(
        "📘 Generar Word"
    ):

        with st.spinner(
            "Generando documentación..."
        ):

            try:

                package_path = extract_package(
                    uploaded_file
                )

                word_file = (
                    generate_word_document(
                        package_path
                    )
                )

                st.download_button(

                    label="⬇️ Descargar Word",

                    data=word_file,

                    file_name=(
                        "documentacion_oic.docx"
                    ),

                    mime=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )

            except Exception as e:

                st.error(
                    f"Error: {str(e)}"
                )


# =========================================================
# UI - IAR
# =========================================================

st.markdown("---")

uploaded_iar = st.file_uploader(
    "Subir archivo .iar",
    type=["iar"]
)

if uploaded_iar:

    st.success(
        "Archivo IAR cargado correctamente."
    )

    if st.button(
        "📦 Descargar contenido interno del IAR"
    ):

        try:

            temp_iar = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=".iar"
            )

            temp_iar.write(
                uploaded_iar.read()
            )

            temp_iar.close()

            extracted_iar = extract_iar(
                temp_iar.name
            )

            zip_content = create_zip_from_folder(
                extracted_iar
            )

            st.download_button(

                label="⬇️ Descargar contenido del IAR",

                data=zip_content,

                file_name="contenido_iar.zip",

                mime="application/zip"
            )

        except Exception as e:

            st.error(
                f"Error extrayendo IAR: {str(e)}"
            )